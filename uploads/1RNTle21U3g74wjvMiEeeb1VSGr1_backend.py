#!/usr/bin/env python3

"""
NOVA PROFESSIONAL API - ULTRA-ENHANCED BACKEND
FastAPI backend with ALL premium features from enhanced_cli.py
Premium Features: Azure Voice, Weather, News, Crypto, Replicate Image Gen, Web Search, Location
Professional focus: English + Hinglish, Enterprise-grade functionality with Firebase Auth
"""

import asyncio
import os
import sys
import json
import time
import threading
import sqlite3
import logging
import hashlib
import re
import requests
import random
import pickle
import base64
import subprocess
import webbrowser
import geocoder
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Union, Tuple
from collections import defaultdict, deque, Counter
from pathlib import Path
import numpy as np

# FastAPI imports
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, BackgroundTasks
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, Field
import uvicorn

# Multi-folder import solution
project_root = os.path.dirname(os.path.abspath(__file__))
folders_to_add = [
    'src',
    os.path.join('src', 'memory'),
    os.path.join('src', 'unique_features'),
    os.path.join('src', 'agents'),
]

for folder in folders_to_add:
    folder_path = os.path.join(project_root, folder)
    if os.path.exists(folder_path) and folder_path not in sys.path:
        sys.path.insert(0, folder_path)

from dotenv import load_dotenv

# Load environment variables
load_dotenv()
FIREBASE_KEY_PATH = os.getenv("FIREBASE_KEY_PATH", "firebase-service-account-key.json")

# ─────────────────────────────────────────
# Firebase Admin SDK (auth verification)
# ─────────────────────────────────────────
try:
    import firebase_admin
    from firebase_admin import credentials, auth
    if not firebase_admin._apps:
        cred = credentials.Certificate(FIREBASE_KEY_PATH)
        firebase_admin.initialize_app(cred)
    FIREBASE_ENABLED = True
except Exception as e:
    FIREBASE_ENABLED = False
    print(f"⚠️ Firebase initialization failed: {e}")

security = HTTPBearer()

async def verify_firebase_token(credentials: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    if not FIREBASE_ENABLED:
        return {"uid": "dev_user", "email": "dev@example.com", "name": "Dev User"}
    
    try:
        decoded_token = auth.verify_id_token(credentials.credentials)
        user = auth.get_user(decoded_token['uid'])
        return {
            "uid": decoded_token['uid'],
            "email": user.email,
            "name": user.display_name or "User",
            "provider": decoded_token.get('firebase', {}).get('sign_in_provider', 'unknown')
        }
    except Exception as e:
        logging.error(f"Firebase auth failed: {e}")
        raise HTTPException(status_code=401, detail="Invalid Firebase token")

# Voice processing imports (Basic + Azure)
try:
    import speech_recognition as sr
    import pyttsx3
    VOICE_AVAILABLE = True
except ImportError:
    VOICE_AVAILABLE = False

# Azure Voice imports (PREMIUM)
try:
    import azure.cognitiveservices.speech as speechsdk
    AZURE_VOICE_AVAILABLE = True
    print("✅ Azure Voice Services loaded!")
except ImportError:
    AZURE_VOICE_AVAILABLE = False
    print("⚠️ Azure Voice not available")

# File processing imports
try:
    from PIL import Image, ImageEnhance, ImageFilter
    import PyPDF2
    import docx
    import openpyxl
    import matplotlib.pyplot as plt
    import seaborn as sns
    import pandas as pd
    import cv2
    FILE_PROCESSING_AVAILABLE = True
    print("✅ File Processing capabilities loaded!")
except ImportError:
    FILE_PROCESSING_AVAILABLE = False
    print("⚠️ File Processing not available")

# Web scraping and search imports (PREMIUM)
try:
    from bs4 import BeautifulSoup
    import feedparser
    WEB_SEARCH_AVAILABLE = True
    print("✅ Web Search capabilities loaded!")
except ImportError:
    WEB_SEARCH_AVAILABLE = False
    print("⚠️ Web Search not available")

# Image generation imports (PREMIUM) - FIXED TO USE REPLICATE
try:
    import replicate
    IMAGE_GENERATION_AVAILABLE = True
    print("✅ Replicate Image Generation loaded!")
except ImportError:
    IMAGE_GENERATION_AVAILABLE = False
    print("⚠️ Image Generation not available")

# GitHub Integration imports - FIXED
try:
    import chromadb
    # Try new langchain imports first with fallback
    try:
        from langchain_community.document_loaders import UnstructuredFileLoader
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        from langchain_community.vectorstores import Chroma
        from langchain_community.embeddings import HuggingFaceEmbeddings
    except ImportError:
        # Fallback to safe text loader if unstructured fails
        try:
            from langchain_community.document_loaders import TextLoader as UnstructuredFileLoader
            from langchain.text_splitter import RecursiveCharacterTextSplitter
            from langchain_community.vectorstores import Chroma
            from langchain_community.embeddings import HuggingFaceEmbeddings
        except ImportError:
            # Final fallback to old imports
            from langchain.document_loaders import TextLoader as UnstructuredFileLoader
            from langchain.text_splitter import RecursiveCharacterTextSplitter
            from langchain.vectorstores import Chroma
            from langchain.embeddings import HuggingFaceEmbeddings
    GITHUB_INTEGRATION = True
    print("✅ GitHub integration loaded!")
except ImportError as e:
    GITHUB_INTEGRATION = False
    print(f"⚠️ GitHub integration not available: {e}")

# Professional Agents Import
try:
    from agents.coding_agent import ProLevelCodingExpert
    from agents.career_coach import ProfessionalCareerCoach
    from agents.business_consultant import SmartBusinessConsultant
    from agents.medical_advisor import SimpleMedicalAdvisor
    from agents.emotional_counselor import SimpleEmotionalCounselor
    from agents.techincal_architect import TechnicalArchitect
    PROFESSIONAL_AGENTS_LOADED = True
    print("✅ Professional agents loaded successfully!")
except ImportError as e:
    PROFESSIONAL_AGENTS_LOADED = False
    print(f"❌ Professional agents import failed: {e}")

# Advanced Systems Import
try:
    from memory.sharp_memory import SharpMemorySystem
    from unique_features.smart_orchestrator import IntelligentAPIOrchestrator
    from unique_features.api_drift_detector import APIPerformanceDrifter
    ADVANCED_SYSTEMS = True
    print("✅ Advanced systems loaded!")
except ImportError as e:
    ADVANCED_SYSTEMS = False
    print(f"⚠️ Advanced systems not available: {e}")

# GitHub Repo Analysis Import - UPDATED
try:
    from agents.ingest import main as ingest_repo, process_and_store_documents
    from agents.qa_engine import create_qa_engine, EnhancedQAEngine
    GITHUB_INTEGRATION = GITHUB_INTEGRATION and True
    print("✅ GitHub QA engine loaded!")
except ImportError as e:
    GITHUB_INTEGRATION = False
    ingest_repo = None
    create_qa_engine = None
    print(f"⚠️ GitHub QA engine not available: {e}")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# FastAPI app initialization
app = FastAPI(
    title="NOVA Ultra Professional API",
    description="Enterprise-grade AI assistant with premium features",
    version="2.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ========== PYDANTIC MODELS ==========
class ChatRequest(BaseModel):
    message: str = Field(..., description="User message")
    session_id: Optional[str] = Field(None, description="Session ID for conversation context")
    voice_mode: Optional[bool] = Field(False, description="Enable voice response")
    language: Optional[str] = Field("en-IN", description="Language preference")

class ChatResponse(BaseModel):
    success: bool
    response: str
    agent_type: str
    agent_confidence: float
    language: str
    emotion: str
    emotion_confidence: float
    response_time: float
    context_used: bool
    location: Optional[str]
    weather_context: Optional[str]
    session_id: str

class VoiceRequest(BaseModel):
    audio_data: Optional[str] = Field(None, description="Base64 encoded audio data")
    language: Optional[str] = Field("en-IN", description="Language for speech recognition")

class VoiceResponse(BaseModel):
    success: bool
    recognized_text: Optional[str]
    response_text: str
    audio_response: Optional[str] = Field(None, description="Base64 encoded audio response")

class WeatherRequest(BaseModel):
    location: Optional[str] = Field(None, description="Location for weather data")

class WeatherResponse(BaseModel):
    success: bool
    location: str
    country: Optional[str]
    temperature: Optional[float]
    feels_like: Optional[float]
    humidity: Optional[int]
    description: Optional[str]
    wind_speed: Optional[float]
    error: Optional[str]

class SearchRequest(BaseModel):
    query: str = Field(..., description="Search query")
    max_results: Optional[int] = Field(5, description="Maximum number of results")
    search_type: Optional[str] = Field("web", description="Type of search: web, news")

class SearchResponse(BaseModel):
    success: bool
    query: str
    results: List[Dict[str, Any]]
    count: int
    error: Optional[str]

class CryptoRequest(BaseModel):
    symbols: Optional[List[str]] = Field(None, description="Cryptocurrency symbols")

class CryptoResponse(BaseModel):
    success: bool
    prices: Dict[str, Any]
    timestamp: str
    error: Optional[str]

class ImageGenRequest(BaseModel):
    prompt: str = Field(..., description="Image generation prompt")
    size: Optional[str] = Field("1024x1024", description="Image size")

class ImageGenResponse(BaseModel):
    success: bool
    prompt: str
    image_url: Optional[str]
    size: str
    model: str
    error: Optional[str]

class GitHubAnalyzeRequest(BaseModel):
    repo_url: str = Field(..., description="GitHub repository URL")

class GitHubAnalyzeResponse(BaseModel):
    success: bool
    repo_name: str
    repo_url: str
    analysis: Dict[str, Any]
    files_processed: Any
    languages: List[str]
    issues_found: List[str]
    suggestions: List[str]
    error: Optional[str]

class FileProcessRequest(BaseModel):
    file_path: str = Field(..., description="Path to file for processing")

class FileProcessResponse(BaseModel):
    success: bool
    file_type: str
    analysis: Dict[str, Any]
    error: Optional[str]

class StatusResponse(BaseModel):
    system_status: Dict[str, Any]
    session_info: Dict[str, Any]
    capabilities: List[Dict[str, str]]

class MemoryStatsResponse(BaseModel):
    total_conversations: int
    recent_conversations: int
    voice_interactions: int
    file_processes: int
    search_queries: int
    agent_distribution: List[Dict[str, Any]]
    emotion_distribution: List[Dict[str, Any]]

# ========== IMPORT ALL CLASSES FROM CLI ==========
# Ultra Hybrid Memory System
class UltraHybridMemorySystem:
    """Ultra Advanced Hybrid Memory with ALL previous features"""
    
    def __init__(self, db_path="nova_ultra_professional_memory.db"):
        # Database setup - FIXED PATH ISSUE
        if not os.path.isabs(db_path):
            self.db_path = os.path.join(os.getcwd(), db_path)
        else:
            self.db_path = db_path
        self.setup_database()
        
        # Memory layers from prev_cli.py (great for conversation flow)
        self.conversation_context = deque(maxlen=100)  # Increased capacity
        self.user_profile = {}
        self.emotional_state = "neutral"
        self.learning_patterns = defaultdict(list)
        self.personality_insights = {}
        self.user_preferences = {}
        self.conversation_history = []
        
        # Memory layers from cli.py (great for technical queries)
        self.short_term_memory = deque(maxlen=200)  # Increased capacity
        self.working_memory = {}
        self.conversation_threads = {}
        self.context_memory = {}
        
        # NEW: Premium memory features
        self.voice_memory = deque(maxlen=50)
        self.file_memory = {}
        self.search_memory = deque(maxlen=30)
        self.image_memory = deque(maxlen=20)
        
        # Semantic memory for technical queries
        self.setup_semantic_memory()
        print("✅ Ultra Hybrid Memory System initialized")

    def setup_database(self):
        """Setup ultra comprehensive database schema"""
        try:
            # Ensure database directory exists
            db_dir = os.path.dirname(self.db_path)
            if db_dir and not os.path.exists(db_dir):
                os.makedirs(db_dir, exist_ok=True)
            
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                # Enhanced conversations table
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS conversations (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id TEXT,
                    session_id TEXT,
                    user_input TEXT,
                    bot_response TEXT,
                    agent_type TEXT,
                    language TEXT,
                    emotion TEXT,
                    confidence REAL,
                    timestamp DATETIME,
                    feedback INTEGER DEFAULT 0,
                    context_summary TEXT,
                    learned_facts TEXT,
                    satisfaction_rating INTEGER,
                    conversation_thread_id TEXT,
                    intent_detected TEXT,
                    response_time REAL,
                    voice_used BOOLEAN DEFAULT 0,
                    location TEXT,
                    weather_context TEXT,
                    search_queries TEXT
                )
                ''')
                
                # Enhanced user profiles
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS user_profiles (
                    user_id TEXT PRIMARY KEY,
                    name TEXT,
                    career_goals TEXT,
                    current_role TEXT,
                    experience_years INTEGER,
                    skills TEXT,
                    preferences TEXT,
                    communication_style TEXT,
                    emotional_patterns TEXT,
                    conversation_patterns TEXT,
                    expertise_level TEXT,
                    topics_of_interest TEXT,
                    last_updated DATETIME,
                    total_conversations INTEGER DEFAULT 0,
                    preferred_voice TEXT,
                    location TEXT,
                    timezone TEXT,
                    personality_type TEXT,
                    learning_style TEXT
                )
                ''')
                
                # GitHub repositories
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS github_repos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    repo_url TEXT UNIQUE,
                    repo_name TEXT,
                    analysis_date DATETIME,
                    file_count INTEGER,
                    languages_detected TEXT,
                    issues_found TEXT,
                    suggestions TEXT,
                    vector_db_path TEXT
                )
                ''')
                
                # Voice interactions
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS voice_interactions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id TEXT,
                    voice_input TEXT,
                    voice_response TEXT,
                    language_detected TEXT,
                    emotion_detected TEXT,
                    voice_engine TEXT,
                    timestamp DATETIME
                )
                ''')
                
                # File processing history
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS file_processing (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id TEXT,
                    file_path TEXT,
                    file_type TEXT,
                    processing_result TEXT,
                    timestamp DATETIME,
                    success BOOLEAN
                )
                ''')
                
                # Search history
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS search_history (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id TEXT,
                    search_query TEXT,
                    search_type TEXT,
                    results_count INTEGER,
                    timestamp DATETIME
                )
                ''')
                
                # Weather queries
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS weather_queries (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id TEXT,
                    location TEXT,
                    weather_data TEXT,
                    timestamp DATETIME
                )
                ''')
                
                conn.commit()
            print("✅ Ultra Database initialized with premium schema")
            
        except Exception as e:
            print(f"⚠️ Database setup error: {e}")

    def setup_semantic_memory(self):
        """Setup semantic memory for technical queries"""
        try:
            if ADVANCED_SYSTEMS:
                self.semantic_memory = SharpMemorySystem()
            else:
                self.semantic_memory = None
        except Exception as e:
            print(f"⚠️ Semantic memory setup error: {e}")
            self.semantic_memory = None

    async def remember_conversation(self, user_id: str, session_id: str,
                                  user_input: str, bot_response: str,
                                  agent_type: str, language: str,
                                  emotion: str, confidence: float,
                                  intent: str = None, response_time: float = 0.0,
                                  voice_used: bool = False, location: str = None,
                                  weather_context: str = None, search_queries: str = None):
        """Ultra enhanced conversation memory storage"""
        try:
            # Extract learning points
            learned_facts = self.extract_learning_points(user_input, bot_response)
            context_summary = self.generate_context_summary()
            
            # Store in database
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                INSERT INTO conversations 
                (user_id, session_id, user_input, bot_response, agent_type,
                 language, emotion, confidence, timestamp, context_summary,
                 learned_facts, conversation_thread_id, intent_detected, response_time,
                 voice_used, location, weather_context, search_queries)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (user_id, session_id, user_input, bot_response, agent_type,
                      language, emotion, confidence, datetime.now(), context_summary,
                      learned_facts, self.generate_thread_id(), intent, response_time,
                      voice_used, location, weather_context, search_queries))
                conn.commit()
            
            # Store in conversation context
            self.conversation_context.append({
                'user': user_input,
                'bot': bot_response,
                'emotion': emotion,
                'agent': agent_type,
                'timestamp': datetime.now(),
                'voice_used': voice_used,
                'location': location
            })
            
            # Store in short-term memory
            memory_entry = {
                'user_id': user_id,
                'session_id': session_id,
                'timestamp': datetime.now().isoformat(),
                'user_input': user_input,
                'ai_response': bot_response,
                'agent_used': agent_type,
                'emotion': emotion,
                'intent': intent,
                'voice_used': voice_used
            }
            self.short_term_memory.append(memory_entry)
            
            # Store in semantic memory for technical queries
            if self.semantic_memory and agent_type in ['coding', 'business', 'technical']:
                try:
                    await self.semantic_memory.remember_conversation_advanced(
                        user_input, bot_response,
                        {'agent_used': agent_type, 'emotion': emotion},
                        user_id, session_id
                    )
                except Exception as e:
                    print(f"⚠️ Semantic memory storage error: {e}")
            
        except Exception as e:
            print(f"⚠️ Memory storage error: {e}")

    def remember_voice_interaction(self, user_id: str, voice_input: str,
                                 voice_response: str, language: str,
                                 emotion: str, voice_engine: str):
        """Remember voice interactions"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                INSERT INTO voice_interactions 
                (user_id, voice_input, voice_response, language_detected,
                 emotion_detected, voice_engine, timestamp)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                ''', (user_id, voice_input, voice_response, language,
                      emotion, voice_engine, datetime.now()))
                conn.commit()
            
            self.voice_memory.append({
                'input': voice_input,
                'response': voice_response,
                'language': language,
                'emotion': emotion,
                'engine': voice_engine,
                'timestamp': datetime.now()
            })
            
        except Exception as e:
            print(f"⚠️ Voice memory error: {e}")

    def remember_file_processing(self, user_id: str, file_path: str,
                               file_type: str, result: str, success: bool):
        """Remember file processing"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                INSERT INTO file_processing 
                (user_id, file_path, file_type, processing_result, timestamp, success)
                VALUES (?, ?, ?, ?, ?, ?)
                ''', (user_id, file_path, file_type, result, datetime.now(), success))
                conn.commit()
            
            self.file_memory[file_path] = {
                'type': file_type,
                'result': result,
                'success': success,
                'timestamp': datetime.now()
            }
            
        except Exception as e:
            print(f"⚠️ File memory error: {e}")

    def remember_search_query(self, user_id: str, query: str,
                            search_type: str, results_count: int):
        """Remember search queries"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                INSERT INTO search_history 
                (user_id, search_query, search_type, results_count, timestamp)
                VALUES (?, ?, ?, ?, ?)
                ''', (user_id, query, search_type, results_count, datetime.now()))
                conn.commit()
            
            self.search_memory.append({
                'query': query,
                'type': search_type,
                'count': results_count,
                'timestamp': datetime.now()
            })
            
        except Exception as e:
            print(f"⚠️ Search memory error: {e}")

    def get_relevant_context(self, user_input: str, user_id: str, limit: int = 15) -> str:
        """Get ultra comprehensive relevant context"""
        try:
            # Get database context
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                SELECT user_input, bot_response, emotion, learned_facts, agent_type,
                       voice_used, location, weather_context
                FROM conversations 
                WHERE user_id = ? 
                ORDER BY timestamp DESC 
                LIMIT ?
                ''', (user_id, limit))
                conversations = cursor.fetchall()
            
            if not conversations:
                return ""
            
            # Build ultra context summary
            context = "Previous conversation context:\n"
            for conv in conversations:
                context += f"[{conv[4].upper()}] User ({conv[2]}): {conv[0][:80]}...\n"
                context += f"NOVA: {conv[1][:80]}...\n"
                if conv[3]:
                    context += f"Learned: {conv[3]}\n"
                if conv[5]:  # voice_used
                    context += f"[VOICE MODE]\n"
                if conv[6]:  # location
                    context += f"Location: {conv[6]}\n"
                if conv[7]:  # weather_context
                    context += f"Weather: {conv[7]}\n"
                context += "---\n"
            
            return context
            
        except Exception as e:
            print(f"⚠️ Context retrieval error: {e}")
            return ""

    def extract_learning_points(self, user_input: str, bot_response: str) -> str:
        """Extract learning points from conversation"""
        learning_keywords = [
            "my name is", "i am", "i work", "i like", "i don't like",
            "my preference", "remember that", "important", "my goal",
            "my project", "my problem", "i need help with", "my role",
            "my company", "my experience", "my skills", "career goal",
            "i live in", "my location", "my city", "my country",
            "i prefer", "i want", "i need", "i use", "my favorite"
        ]
        
        learned = []
        user_lower = user_input.lower()
        
        for keyword in learning_keywords:
            if keyword in user_lower:
                sentences = user_input.split('.')
                for sentence in sentences:
                    if keyword in sentence.lower():
                        learned.append(sentence.strip())
        
        return "; ".join(learned)

    def generate_context_summary(self) -> str:
        """Generate ultra context summary from recent conversations"""
        if not self.conversation_context:
            return ""
        
        recent_topics = []
        emotions = []
        agents = []
        voice_usage = []
        locations = []
        
        for conv in list(self.conversation_context)[-10:]:
            recent_topics.append(conv['user'][:50])
            emotions.append(conv['emotion'])
            agents.append(conv['agent'])
            if conv.get('voice_used'):
                voice_usage.append(True)
            if conv.get('location'):
                locations.append(conv['location'])
        
        dominant_emotion = max(set(emotions), key=emotions.count) if emotions else "neutral"
        most_used_agent = max(set(agents), key=agents.count) if agents else "general"
        voice_percentage = (len(voice_usage) / len(emotions)) * 100 if emotions else 0
        
        summary = f"Recent topics: {'; '.join(recent_topics)}. "
        summary += f"Emotion: {dominant_emotion}. Agent: {most_used_agent}. "
        if voice_percentage > 0:
            summary += f"Voice usage: {voice_percentage:.0f}%. "
        if locations:
            summary += f"Location context: {locations[-1]}."
        
        return summary

    def generate_thread_id(self) -> str:
        """Generate conversation thread ID"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return f"thread_{timestamp}_{random.randint(1000, 9999)}"

    def get_memory_stats(self, user_id: str) -> Dict[str, Any]:
        """Get comprehensive memory statistics"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                # Total conversations
                cursor.execute('SELECT COUNT(*) FROM conversations WHERE user_id = ?', (user_id,))
                total_conversations = cursor.fetchone()[0]
                
                # Recent conversations
                week_ago = (datetime.now() - timedelta(days=7)).isoformat()
                cursor.execute('SELECT COUNT(*) FROM conversations WHERE user_id = ? AND timestamp > ?', (user_id, week_ago))
                recent_conversations = cursor.fetchone()[0]
                
                # Voice interactions
                cursor.execute('SELECT COUNT(*) FROM voice_interactions WHERE user_id = ?', (user_id,))
                voice_interactions = cursor.fetchone()[0]
                
                # File processing
                cursor.execute('SELECT COUNT(*) FROM file_processing WHERE user_id = ?', (user_id,))
                file_processes = cursor.fetchone()[0]
                
                # Search history  
                cursor.execute('SELECT COUNT(*) FROM search_history WHERE user_id = ?', (user_id,))
                search_queries = cursor.fetchone()[0]
                
                # Agent distribution
                cursor.execute('''
                SELECT agent_type, COUNT(*) as count 
                FROM conversations 
                WHERE user_id = ?
                GROUP BY agent_type 
                ORDER BY count DESC
                ''', (user_id,))
                agent_stats = cursor.fetchall()
                
                # Emotion distribution
                cursor.execute('''
                SELECT emotion, COUNT(*) as count 
                FROM conversations 
                WHERE user_id = ?
                GROUP BY emotion 
                ORDER BY count DESC 
                LIMIT 5
                ''', (user_id,))
                emotion_stats = cursor.fetchall()
                
                return {
                    "total_conversations": total_conversations,
                    "recent_conversations": recent_conversations,
                    "voice_interactions": voice_interactions,
                    "file_processes": file_processes,
                    "search_queries": search_queries,
                    "agent_distribution": [{"agent": agent, "count": count} for agent, count in agent_stats],
                    "emotion_distribution": [{"emotion": emotion, "count": count} for emotion, count in emotion_stats],
                    "short_term_memory": len(self.short_term_memory),
                    "conversation_context": len(self.conversation_context),
                    "voice_memory": len(self.voice_memory),
                    "search_memory": len(self.search_memory),
                    "image_memory": len(self.image_memory)
                }
                
        except Exception as e:
            logger.error(f"Memory stats error: {e}")
            return {}

# Premium Azure Voice System
class PremiumAzureVoiceSystem:
    """Premium Azure Voice System with multi-language support"""
    
    def __init__(self):
        self.azure_enabled = AZURE_VOICE_AVAILABLE
        self.basic_voice_enabled = VOICE_AVAILABLE
        self.speech_config = None
        self.speech_recognizer = None
        self.speech_synthesizer = None
        
        # Basic voice fallback
        self.tts_engine = None
        self.recognizer = None
        
        # Initialize Azure Voice (Premium)
        if self.azure_enabled:
            try:
                self.setup_azure_voice()
                print("✅ Premium Azure Voice System initialized")
            except Exception as e:
                print(f"⚠️ Azure Voice setup error: {e}")
                self.azure_enabled = False
        
        # Initialize Basic Voice (Fallback)
        if self.basic_voice_enabled:
            try:
                self.setup_basic_voice()
                print("✅ Basic Voice System initialized as fallback")
            except Exception as e:
                self.basic_voice_enabled = False
                print(f"⚠️ Basic voice setup error: {e}")

    def setup_azure_voice(self):
        """Setup premium Azure voice services"""
        azure_key = os.getenv('AZURE_SPEECH_KEY')
        azure_region = os.getenv('AZURE_SPEECH_REGION', 'eastus')
        
        if not azure_key:
            raise Exception("Azure Speech Key not found")
        
        self.speech_config = speechsdk.SpeechConfig(
            subscription=azure_key,
            region=azure_region
        )
        
        # Configure for Indian English with professional voice
        self.speech_config.speech_recognition_language = "en-IN"
        self.speech_config.speech_synthesis_voice_name = "en-IN-NeerjaNeural"
        
        # Create recognizer and synthesizer
        audio_config = speechsdk.audio.AudioConfig(use_default_microphone=True)
        self.speech_recognizer = speechsdk.SpeechRecognizer(
            speech_config=self.speech_config,
            audio_config=audio_config
        )
        
        self.speech_synthesizer = speechsdk.SpeechSynthesizer(
            speech_config=self.speech_config
        )

    def setup_basic_voice(self):
        """Setup basic voice as fallback"""
        self.recognizer = sr.Recognizer()
        self.tts_engine = pyttsx3.init()
        
        # Configure TTS for professional quality
        voices = self.tts_engine.getProperty('voices')
        if voices:
            for voice in voices:
                if 'female' in voice.name.lower() or 'zira' in voice.name.lower():
                    self.tts_engine.setProperty('voice', voice.id)
                    break
        
        self.tts_engine.setProperty('rate', 180)
        self.tts_engine.setProperty('volume', 0.9)

    async def recognize_from_audio_data(self, audio_data: str, language: str = "en-IN") -> Optional[str]:
        """Recognize speech from base64 audio data"""
        try:
            # Decode base64 audio data
            audio_bytes = base64.b64decode(audio_data)
            
            # For now, return a placeholder - in production, implement actual audio processing
            return "Speech recognition from audio data not yet implemented"
            
        except Exception as e:
            logger.error(f"Audio recognition error: {e}")
            return None

    async def synthesize_speech(self, text: str, language: str = "en-IN", voice: str = None) -> Optional[str]:
        """Synthesize speech and return base64 encoded audio"""
        if not self.azure_enabled:
            return None
        
        try:
            # Clean text for TTS
            clean_text = self.clean_text_for_tts(text)
            if not clean_text:
                return None
            
            # Set voice based on language
            voice_map = {
                "en-IN": "en-IN-NeerjaNeural",
                "en-US": "en-US-AriaNeural",
                "hi-IN": "hi-IN-SwaraNeural",
                "en-GB": "en-GB-SoniaNeural"
            }
            
            selected_voice = voice or voice_map.get(language, "en-IN-NeerjaNeural")
            self.speech_config.speech_synthesis_voice_name = selected_voice
            
            # Synthesize speech
            result = await asyncio.to_thread(
                self.speech_synthesizer.speak_text, clean_text
            )
            
            if result.reason == speechsdk.ResultReason.SynthesizingAudioCompleted:
                # Convert audio data to base64
                audio_data = result.audio_data
                return base64.b64encode(audio_data).decode('utf-8')
            else:
                logger.error(f"Azure TTS error: {result.reason}")
                return None
                    
        except Exception as e:
            logger.error(f"Azure TTS error: {e}")
            return None

    def clean_text_for_tts(self, text: str) -> str:
        """Clean text for professional TTS"""
        if isinstance(text, dict):
            text = text.get('content', str(text))
        
        # Remove markdown and formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)     # Italic
        text = re.sub(r'`(.*?)`', r'\1', text)       # Code
        text = re.sub(r'#{1,6}\s*(.*?)(?:\n|$)', r'\1', text)  # Headers
        text = re.sub(r'\[.*?\]\(.*?\)', '', text)   # Links
        
        # Remove emojis and special characters
        text = re.sub(r'[🔧💼📈🏥💙🚀🎯📋💡📚🤖⚠️✅❌🔊📝🎤🌤️🌧️⛈️🌙☀️🌍📍💰🔍📰🖼️]', '', text)
        
        # Limit length for professional delivery
        if len(text) > 500:
            text = text[:500] + "... Please check the full response on screen."
        
        return text.strip()

    def get_available_voices(self, language: str = "en-IN") -> List[str]:
        """Get available voices for language"""
        voice_options = {
            "en-IN": ["en-IN-NeerjaNeural", "en-IN-PrabhatNeural"],
            "en-US": ["en-US-AriaNeural", "en-US-DavisNeural", "en-US-JennyNeural"],
            "hi-IN": ["hi-IN-SwaraNeural", "hi-IN-MadhurNeural"],
            "en-GB": ["en-GB-SoniaNeural", "en-GB-RyanNeural"]
        }
        
        return voice_options.get(language, ["en-IN-NeerjaNeural"])

# Premium Web Search System
class PremiumWebSearchSystem:
    """Premium web search with DuckDuckGo scraping (No API needed)"""
    
    def __init__(self):
        self.search_enabled = True  # Always enabled now
        print("✅ Premium Web Search System initialized (Free DuckDuckGo)")

    def extract_domain(self, url: str) -> str:
        """Extract domain from URL"""
        try:
            from urllib.parse import urlparse
            return urlparse(url).netloc
        except:
            return "Unknown"

    async def search_web(self, query: str, max_results: int = 5) -> Dict[str, Any]:
        """Free DuckDuckGo web search via scraping"""
        try:
            logger.info(f"Searching web for: {query}")
            
            # DuckDuckGo HTML search URL
            url = f"https://html.duckduckgo.com/html/?q={query}&kl=in-en"
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            results = []
            
            # Parse DuckDuckGo results
            for result_div in soup.find_all('div', class_='result__body')[:max_results]:
                try:
                    title_elem = result_div.find('a', class_='result__a')
                    snippet_elem = result_div.find('a', class_='result__snippet')
                    
                    if title_elem:
                        title = title_elem.get_text().strip()
                        url_link = title_elem.get('href', '')
                        snippet = snippet_elem.get_text().strip() if snippet_elem else "No description"
                        
                        results.append({
                            "title": title,
                            "snippet": snippet,
                            "url": url_link,
                            "source": self.extract_domain(url_link)
                        })
                except Exception as e:
                    continue
            
            return {
                "success": True,
                "query": query,
                "results": results,
                "count": len(results)
            }
            
        except Exception as e:
            return {"error": f"Web search failed: {e}"}

    async def search_news(self, query: str = "latest news", max_results: int = 5) -> Dict[str, Any]:
        """Free DuckDuckGo news search"""
        try:
            logger.info(f"Searching news for: {query}")
            
            # DuckDuckGo news search
            news_query = f"{query} news"
            url = f"https://html.duckduckgo.com/html/?q={news_query}&iar=news&kl=in-en"
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            results = []
            
            for result_div in soup.find_all('div', class_='result__body')[:max_results]:
                try:
                    title_elem = result_div.find('a', class_='result__a')
                    snippet_elem = result_div.find('a', class_='result__snippet')
                    
                    if title_elem:
                        title = title_elem.get_text().strip()
                        url_link = title_elem.get('href', '')
                        snippet = snippet_elem.get_text().strip() if snippet_elem else "No description"
                        
                        results.append({
                            "title": title,
                            "snippet": snippet,
                            "url": url_link,
                            "source": self.extract_domain(url_link),
                            "date": "Recent"  # DuckDuckGo doesn't provide exact dates
                        })
                except Exception as e:
                    continue
            
            return {
                "success": True,
                "query": query,
                "results": results,
                "count": len(results)
            }
            
        except Exception as e:
            return {"error": f"News search failed: {e}"}

# Premium Weather System
class PremiumWeatherSystem:
    """Premium weather system with location detection"""
    
    def __init__(self):
        self.weather_api_key = os.getenv('OPENWEATHER_API_KEY')
        self.weather_enabled = bool(self.weather_api_key)
        if self.weather_enabled:
            print("✅ Premium Weather System initialized")

    async def get_weather(self, location: str = None) -> Dict[str, Any]:
        """Get weather information"""
        if not self.weather_enabled:
            return {"error": "Weather API key not configured"}
        
        try:
            # Auto-detect location if not provided
            if not location:
                location = self.get_current_location()
            
            logger.info(f"Getting weather for: {location}")
            
            # Get weather from OpenWeatherMap
            url = f"http://api.openweathermap.org/data/2.5/weather"
            params = {
                'q': location,
                'appid': self.weather_api_key,
                'units': 'metric'
            }
            
            response = requests.get(url, params=params, timeout=10)
            response.raise_for_status()
            data = response.json()
            
            return {
                "success": True,
                "location": data['name'],
                "country": data['sys']['country'],
                "temperature": data['main']['temp'],
                "feels_like": data['main']['feels_like'],
                "humidity": data['main']['humidity'],
                "pressure": data['main']['pressure'],
                "description": data['weather'][0]['description'],
                "icon": data['weather'][0]['icon'],
                "wind_speed": data['wind']['speed'],
                "visibility": data.get('visibility', 'N/A')
            }
            
        except Exception as e:
            return {"error": f"Weather fetch failed: {e}"}

    def get_current_location(self) -> str:
        """Auto-detect current location"""
        try:
            g = geocoder.ip('me')
            if g.ok:
                return f"{g.city}, {g.country}"
            else:
                return "Mumbai, India"  # Default fallback
        except:
            return "Mumbai, India"  # Default fallback

    async def get_forecast(self, location: str = None, days: int = 5) -> Dict[str, Any]:
        """Get weather forecast"""
        if not self.weather_enabled:
            return {"error": "Weather API key not configured"}
        
        try:
            if not location:
                location = self.get_current_location()
            
            url = f"http://api.openweathermap.org/data/2.5/forecast"
            params = {
                'q': location,
                'appid': self.weather_api_key,
                'units': 'metric',
                'cnt': days * 8  # 3-hour intervals
            }
            
            response = requests.get(url, params=params, timeout=10)
            response.raise_for_status()
            data = response.json()
            
            forecast = []
            for item in data['list']:
                forecast.append({
                    "datetime": item['dt_txt'],
                    "temperature": item['main']['temp'],
                    "description": item['weather'][0]['description'],
                    "humidity": item['main']['humidity']
                })
            
            return {
                "success": True,
                "location": data['city']['name'],
                "forecast": forecast
            }
            
        except Exception as e:
            return {"error": f"Forecast fetch failed: {e}"}

# Premium Crypto System
class PremiumCryptoSystem:
    """Premium cryptocurrency tracking system"""
    
    def __init__(self):
        self.crypto_enabled = True
        print("✅ Premium Crypto System initialized")

    async def get_crypto_prices(self, symbols: List[str] = None) -> Dict[str, Any]:
        """Get cryptocurrency prices"""
        if symbols is None:
            symbols = ['bitcoin', 'ethereum', 'binancecoin', 'cardano', 'solana']
        
        try:
            logger.info(f"Getting crypto prices...")
            
            # Use CoinGecko API (free)
            url = "https://api.coingecko.com/api/v3/simple/price"
            params = {
                'ids': ','.join(symbols),
                'vs_currencies': 'usd,inr',
                'include_24hr_change': 'true',
                'include_market_cap': 'true'
            }
            
            response = requests.get(url, params=params, timeout=10)
            response.raise_for_status()
            data = response.json()
            
            prices = {}
            for symbol, info in data.items():
                prices[symbol] = {
                    "usd": info.get('usd', 0),
                    "inr": info.get('inr', 0),
                    "change_24h": info.get('usd_24h_change', 0),
                    "market_cap": info.get('usd_market_cap', 0)
                }
            
            return {
                "success": True,
                "prices": prices,
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {"error": f"Crypto price fetch failed: {e}"}

    async def get_trending_cryptos(self) -> Dict[str, Any]:
        """Get trending cryptocurrencies"""
        try:
            url = "https://api.coingecko.com/api/v3/search/trending"
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            data = response.json()
            
            trending = []
            for coin in data['coins']:
                trending.append({
                    "name": coin['item']['name'],
                    "symbol": coin['item']['symbol'],
                    "market_cap_rank": coin['item']['market_cap_rank'],
                    "price_btc": coin['item']['price_btc']
                })
            
            return {
                "success": True,
                "trending": trending
            }
            
        except Exception as e:
            return {"error": f"Trending crypto fetch failed: {e}"}

# Premium Image Generation System
class PremiumImageGenerationSystem:
    """Premium image generation using Replicate (Free Models)"""
    
    def __init__(self):
        self.replicate_token = os.getenv('REPLICATE_API_TOKEN')
        self.image_gen_enabled = IMAGE_GENERATION_AVAILABLE and bool(self.replicate_token)
        
        if self.image_gen_enabled:
            try:
                replicate.Client(api_token=self.replicate_token)
                print("✅ Replicate Image Generation System initialized")
            except Exception as e:
                print(f"⚠️ Replicate initialization error: {e}")
                self.image_gen_enabled = False
        else:
            print("⚠️ Replicate token not found, image generation disabled")

    async def generate_image(self, prompt: str, size: str = "1024x1024") -> Dict[str, Any]:
        """Generate image using free Replicate SDXL model"""
        if not self.image_gen_enabled:
            return {"error": "Image generation not available"}
        
        try:
            logger.info(f"Generating image with Replicate SDXL: {prompt}")
            
            # Use free SDXL model
            output = await asyncio.to_thread(
                replicate.run,
                "stability-ai/sdxl",
                input={
                    "prompt": prompt,
                    "width": 1024,
                    "height": 1024,
                    "num_outputs": 1,
                    "scheduler": "K_EULER",
                    "num_inference_steps": 20,
                    "guidance_scale": 7.5
                }
            )
            
            # Extract image URL
            image_url = output[0] if output and len(output) > 0 else None
            
            if image_url:
                return {
                    "success": True,
                    "prompt": prompt,
                    "image_url": image_url,
                    "size": size,
                    "model": "SDXL (Replicate)"
                }
            else:
                return {"error": "Replicate returned no image URL"}
                
        except Exception as e:
            return {"error": f"Image generation failed: {e}"}

    async def edit_image(self, image_path: str, prompt: str) -> Dict[str, Any]:
        """Edit image using Replicate (if model supports it)"""
        if not self.image_gen_enabled:
            return {"error": "Image editing not available"}
        
        try:
            # This would require a specific image editing model
            # For now, return a helpful message
            return {
                "error": "Image editing not yet implemented with current Replicate models",
                "suggestion": "Try generating a new image with your requirements"
            }
            
        except Exception as e:
            return {"error": f"Image editing failed: {e}"}

# GitHub Repo Analyzer
class GitHubRepoAnalyzer:
    """Advanced GitHub repository analyzer with bug detection"""
    
    def __init__(self):
        self.active_repo = None
        self.repo_data = {}
        self.qa_engine = None
        self.vector_db_path = None
        
        if GITHUB_INTEGRATION and create_qa_engine:
            try:
                self.qa_engine = create_qa_engine(simple=False)
                print("✅ GitHub QA Engine initialized")
            except Exception as e:
                try:
                    self.qa_engine = create_qa_engine(simple=True)
                    print("⚠️ Using simple QA Engine")
                except Exception as e2:
                    print(f"⚠️ QA Engine initialization failed: {e2}")
                    self.qa_engine = None

    async def analyze_repository(self, repo_url: str) -> Dict[str, Any]:
        """Analyze GitHub repository comprehensively"""
        if not GITHUB_INTEGRATION or not ingest_repo:
            return {"error": "GitHub integration not available"}
        
        try:
            logger.info(f"Analyzing repository: {repo_url}")
            
            # Extract repo info
            repo_name = repo_url.split('/')[-1].replace('.git', '')
            
            # Use ingest.py to clone and process
            try:
                ingest_repo(repo_url)
                print("✅ Repository ingested successfully")
            except Exception as e:
                return {"error": f"Failed to ingest repository: {e}"}
            
            # Store repo information
            self.active_repo = repo_url
            self.repo_data = {
                'name': repo_name,
                'url': repo_url,
                'analyzed_at': datetime.now(),
                'vector_db_path': "./chroma_db"
            }
            
            # Analyze repository structure
            analysis = await self.perform_code_analysis()
            
            return {
                "success": True,
                "repo_name": repo_name,
                "repo_url": repo_url,
                "analysis": analysis,
                "files_processed": analysis.get('file_count', 0),
                "languages": analysis.get('languages', []),
                "issues_found": analysis.get('issues', []),
                "suggestions": analysis.get('suggestions', [])
            }
            
        except Exception as e:
            return {"error": f"Repository analysis failed: {e}"}

    async def perform_code_analysis(self) -> Dict[str, Any]:
        """Perform comprehensive code analysis"""
        if not self.qa_engine:
            return {
                "error": "QA engine not available",
                'file_count': 'Repository processed',
                'languages': ['Python', 'JavaScript', 'Other'],
                'issues': ["Analysis engine unavailable"],
                'suggestions': ["Manual code review recommended"],
                'detailed_analysis': {}
            }
        
        # Basic analysis questions
        analysis_questions = [
            "What is the main purpose of this codebase?",
            "What programming languages are used?",
            "Are there any potential bugs or issues in the code?",
            "What improvements can be made to this code?",
            "What is the overall structure and architecture?"
        ]
        
        analysis_results = {}
        for question in analysis_questions:
            try:
                result = self.qa_engine.ask(question)
                if isinstance(result, dict) and 'response' in result:
                    analysis_results[question] = result['response']
                else:
                    analysis_results[question] = str(result)
            except Exception as e:
                analysis_results[question] = f"Analysis failed: {e}"
        
        # Extract structured information
        return {
            'file_count': 'Multiple files processed',
            'languages': ['Python', 'JavaScript', 'Other'],  # Simplified
            'issues': self.extract_issues(analysis_results),
            'suggestions': self.extract_suggestions(analysis_results),
            'detailed_analysis': analysis_results
        }

    def extract_issues(self, analysis: Dict[str, str]) -> List[str]:
        """Extract potential issues from analysis"""
        issues = []
        bug_question = "Are there any potential bugs or issues in the code?"
        
        if bug_question in analysis:
            bug_analysis = analysis[bug_question].lower()
            if 'bug' in bug_analysis or 'issue' in bug_analysis or 'error' in bug_analysis:
                issues.append("Potential bugs detected in codebase")
            if 'security' in bug_analysis:
                issues.append("Security concerns identified")
            if 'performance' in bug_analysis:
                issues.append("Performance optimizations needed")
        
        return issues if issues else ["No critical issues detected"]

    def extract_suggestions(self, analysis: Dict[str, str]) -> List[str]:
        """Extract improvement suggestions"""
        suggestions = []
        improvement_question = "What improvements can be made to this code?"
        
        if improvement_question in analysis:
            suggestions.append("Code structure and architecture improvements")
            suggestions.append("Documentation and comments enhancement")
            suggestions.append("Error handling and validation improvements")
            suggestions.append("Performance optimization opportunities")
        
        return suggestions

    async def answer_repo_question(self, question: str) -> str:
        """Answer questions about the active repository"""
        if not self.active_repo or not self.qa_engine:
            return "No active repository or QA engine not available"
        
        try:
            result = self.qa_engine.ask(question)
            if isinstance(result, dict) and 'response' in result:
                return result['response']
            return str(result)
        except Exception as e:
            return f"Failed to answer repository question: {e}"

    def has_active_repo(self) -> bool:
        """Check if there's an active repository"""
        return self.active_repo is not None

# Professional Language Detector
class ProfessionalLanguageDetector:
    """Professional language detection: English + Hinglish only"""
    
    def __init__(self):
        self.language_patterns = {
            "english": [
                "what", "how", "when", "where", "why", "can", "will",
                "should", "could", "would", "the", "and", "but", "this",
                "that", "good", "bad", "right", "wrong", "please", "thank",
                "help", "need", "want", "like", "make", "work", "time"
            ],
            "hinglish": [
                "yaar", "bhai", "dude", "boss", "sir", "madam", "ji",
                "na", "haan", "nahi", "accha", "theek", "kya", "hai",
                "hoon", "main", "tum", "aur", "kar", "kaise", "kyun",
                "matlab", "samjha", "pata", "chal"
            ]
        }
        
        self.hindi_indicators = [
            "hai", "hoon", "kya", "aur", "main", "tum", "yeh", "woh",
            "kaise", "kab", "kahan", "kyun", "matlab", "samjha", "pata"
        ]

    def detect_language(self, text: str) -> str:
        """Detect if text is English or Hinglish"""
        text_words = text.lower().split()
        
        english_count = sum(1 for word in self.language_patterns["english"] if word in text_words)
        hinglish_count = sum(1 for word in self.language_patterns["hinglish"] if word in text_words)
        hindi_count = sum(1 for word in self.hindi_indicators if word in text_words)
        
        total_words = len(text_words)
        if total_words == 0:
            return "english"
        
        # If significant Hindi/Hinglish words detected
        if (hinglish_count + hindi_count) / total_words > 0.15:
            return "hinglish"
        
        return "english"

# Advanced Emotion Detector
class AdvancedEmotionDetector:
    """PhD-level emotion detection with confidence scoring"""
    
    def __init__(self):
        self.emotion_patterns = {
            "excited": [
                "excited", "amazing", "awesome", "fantastic", "great",
                "wonderful", "thrilled", "happy", "joy", "ecstatic",
                "pumped", "energetic", "love", "brilliant", "perfect"
            ],
            "frustrated": [
                "frustrated", "annoyed", "irritated", "angry", "mad",
                "upset", "pissed", "fed up", "bothered", "stressed",
                "hate", "terrible", "awful", "worst", "stupid"
            ],
            "sad": [
                "sad", "depressed", "down", "blue", "unhappy", "miserable",
                "heartbroken", "grief", "disappointed", "dejected",
                "lonely", "empty", "hopeless", "crying", "tears"
            ],
            "anxious": [
                "anxious", "worried", "nervous", "stressed", "panic",
                "fear", "scared", "concern", "overwhelmed", "tense",
                "pressure", "burden", "difficult", "problem", "issue"
            ],
            "confident": [
                "confident", "sure", "certain", "positive", "optimistic",
                "determined", "ready", "motivated", "strong", "capable",
                "can do", "will do", "believe", "achieve", "success"
            ],
            "confused": [
                "confused", "lost", "unclear", "puzzled", "bewildered",
                "don't understand", "help me understand", "not sure",
                "complicated", "difficult", "hard", "stuck", "doubt"
            ],
            "curious": [
                "curious", "interesting", "wonder", "explore", "discover",
                "learn", "understand", "find out", "know more", "research"
            ]
        }

    def detect_emotion(self, text: str) -> Tuple[str, float]:
        """Detect emotion with confidence score"""
        text_lower = text.lower()
        emotion_scores = {}
        
        for emotion, keywords in self.emotion_patterns.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score > 0:
                emotion_scores[emotion] = score
        
        if not emotion_scores:
            return "neutral", 0.5
        
        detected_emotion = max(emotion_scores, key=emotion_scores.get)
        confidence = emotion_scores[detected_emotion] / len(text_lower.split())
        
        return detected_emotion, min(confidence * 3, 1.0)  # Amplify confidence

# Smart API Manager
class SmartAPIManager:
    """Enhanced API manager with professional providers"""
    
    def __init__(self):
        self.providers = [
            {
                "name": "Groq",
                "url": "https://api.groq.com/openai/v1/chat/completions",
                "models": [
                    "llama-3.1-8b-instant",     # fast + free
                    "llama-3.1-70b-versatile",  # higher quality
                    "mixtral-8x7b-32768"        # alternative
                ],
                "headers": lambda: {
                    "Authorization": f"Bearer {os.getenv('GROQ_API_KEY', '')}",
                    "Content-Type": "application/json"
                }
            },
            {
                "name": "OpenRouter",
                "url": "https://openrouter.ai/api/v1/chat/completions",
                "models": [
                    "mistralai/mistral-7b-instruct:free",
                    "meta-llama/llama-3.1-70b-instruct",
                    "anthropic/claude-3.5-sonnet",
                    "google/gemini-pro-1.5"
                ],
                "headers": lambda: {
                    "Authorization": f"Bearer {os.getenv('OPENROUTER_API_KEY', '')}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://nova-professional.ai",
                    "X-Title": "NOVA Professional"
                }
            }
        ]
        
        self.available = []
        for provider in self.providers:
            key_name = f"{provider['name'].upper()}_API_KEY"
            if os.getenv(key_name):
                self.available.append(provider)
        
        self.current = self.available[0] if self.available else None

    async def get_ai_response(self, user_input: str, system_prompt: str,
                            model_preference: str = None) -> Optional[str]:
        """Get AI response with fallback handling"""
        if not self.current:
            return None
        
        # Select model based on preference
        models_to_try = self.current["models"]
        if model_preference:
            preferred_models = [m for m in models_to_try if model_preference in m.lower()]
            models_to_try = preferred_models + [m for m in models_to_try if m not in preferred_models]
        
        for model in models_to_try[:2]:  # Try top 2 models
            try:
                response = requests.post(
                    self.current["url"],
                    headers=self.current["headers"](),
                    json={
                        "model": model,
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_input}
                        ],
                        "max_tokens": 1500,
                        "temperature": 0.7,
                        "top_p": 0.9
                    },
                    timeout=30
                )
                
                if response.status_code == 200:
                    result = response.json()
                    if 'choices' in result and len(result['choices']) > 0:
                        return result['choices'][0]['message']['content'].strip()
                        
            except Exception as e:
                logger.error(f"Model {model} failed: {e}")
                continue
        
        return None

# Ultra File Processor
class UltraFileProcessor:
    """Ultra file processing system with all formats"""
    
    def __init__(self):
        self.file_processing_enabled = FILE_PROCESSING_AVAILABLE

    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Process various file types professionally"""
        if not self.file_processing_enabled:
            return {"error": "File processing not available"}
        
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return {"error": "File not found"}
            
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.pdf':
                return self.process_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self.process_word(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                return self.process_excel(file_path)
            elif file_extension in ['.csv']:
                return self.process_csv(file_path)
            elif file_extension in ['.txt', '.md']:
                return self.process_text(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp']:
                return self.process_image(file_path)
            elif file_extension in ['.py', '.js', '.html', '.css', '.java', '.cpp']:
                return self.process_code(file_path)
            else:
                return {"error": f"Unsupported file type: {file_extension}"}
                
        except Exception as e:
            return {"error": f"File processing failed: {e}"}

    def process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
            
            return {
                "file_type": "pdf",
                "pages": len(pdf_reader.pages),
                "text": text,
                "word_count": len(text.split()),
                "char_count": len(text),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"PDF processing failed: {e}"}

    def process_word(self, file_path: Path) -> Dict[str, Any]:
        """Process Word files"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "file_type": "word",
                "text": text,
                "word_count": len(text.split()),
                "paragraph_count": len(doc.paragraphs),
                "char_count": len(text),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"Word processing failed: {e}"}

    def process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Process Excel files"""
        try:
            df = pd.read_excel(file_path)
            
            return {
                "file_type": "excel",
                "shape": df.shape,
                "columns": list(df.columns),
                "data_types": df.dtypes.to_dict(),
                "summary": df.describe().to_dict(),
                "null_values": df.isnull().sum().to_dict(),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"Excel processing failed: {e}"}

    def process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            df = pd.read_csv(file_path)
            
            return {
                "file_type": "csv",
                "shape": df.shape,
                "columns": list(df.columns),
                "data_types": df.dtypes.to_dict(),
                "summary": df.describe().to_dict(),
                "sample_data": df.head().to_dict(),
                "null_values": df.isnull().sum().to_dict(),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"CSV processing failed: {e}"}

    def process_text(self, file_path: Path) -> Dict[str, Any]:
        """Process text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            words = text.split()
            lines = text.split('\n')
            
            return {
                "file_type": "text",
                "text": text,
                "word_count": len(words),
                "line_count": len(lines),
                "char_count": len(text),
                "average_words_per_line": len(words) / len(lines) if lines else 0,
                "success": True
            }
            
        except Exception as e:
            return {"error": f"Text processing failed: {e}"}

    def process_image(self, file_path: Path) -> Dict[str, Any]:
        """Process image files"""
        try:
            img = Image.open(file_path)
            
            return {
                "file_type": "image",
                "format": img.format,
                "mode": img.mode,
                "size": img.size,
                "width": img.width,
                "height": img.height,
                "has_transparency": img.mode in ('RGBA', 'LA'),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"Image processing failed: {e}"}

    def process_code(self, file_path: Path) -> Dict[str, Any]:
        """Process code files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                code = file.read()
            
            lines = code.split('\n')
            code_lines = [line for line in lines if line.strip() and not line.strip().startswith('#')]
            comment_lines = [line for line in lines if line.strip().startswith('#')]
            
            return {
                "file_type": "code",
                "language": file_path.suffix[1:],
                "total_lines": len(lines),
                "code_lines": len(code_lines),
                "comment_lines": len(comment_lines),
                "blank_lines": len(lines) - len(code_lines) - len(comment_lines),
                "char_count": len(code),
                "functions": self.count_functions(code, file_path.suffix),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"Code processing failed: {e}"}

    def count_functions(self, code: str, extension: str) -> int:
        """Count functions in code"""
        try:
            if extension == '.py':
                return len(re.findall(r'def\s+\w+', code))
            elif extension == '.js':
                return len(re.findall(r'function\s+\w+', code))
            elif extension in ['.java', '.cpp', '.c']:
                return len(re.findall(r'\w+\s+\w+\s*\([^)]*\)\s*{', code))
            else:
                return 0
        except:
            return 0

# ========== MAIN NOVA BACKEND CLASS ==========
class NovaUltraProfessionalBackend:
    """NOVA Ultra Professional Backend - Complete Integrated System with ALL features"""
    
    def __init__(self):
        """Initialize the complete ultra professional system"""
        # Core systems
        self.memory = UltraHybridMemorySystem()
        self.language_detector = ProfessionalLanguageDetector()
        self.emotion_detector = AdvancedEmotionDetector()
        self.api_manager = SmartAPIManager()
        
        # Premium systems
        self.voice_system = PremiumAzureVoiceSystem()
        self.web_search = PremiumWebSearchSystem()
        self.weather_system = PremiumWeatherSystem()
        self.crypto_system = PremiumCryptoSystem()
        self.image_generator = PremiumImageGenerationSystem()
        self.file_processor = UltraFileProcessor()
        self.github_analyzer = GitHubRepoAnalyzer()
        
        # Load professional agents
        self.agents = {}
        if PROFESSIONAL_AGENTS_LOADED:
            try:
                self.agents = {
                    'coding': ProLevelCodingExpert(),
                    'career': ProfessionalCareerCoach(),
                    'business': SmartBusinessConsultant(),
                    'medical': SimpleMedicalAdvisor(),
                    'emotional': SimpleEmotionalCounselor(),
                    'technical_architect': TechnicalArchitect()
                }
                print("✅ Professional agents loaded")
            except Exception as e:
                print(f"⚠️ Agent loading error: {e}")
        
        # Advanced systems
        self.orchestrator = None
        self.drift_detector = None
        if ADVANCED_SYSTEMS:
            try:
                self.smart_orchestrator = IntelligentAPIOrchestrator()
                self.api_drift_detector = APIPerformanceDrifter()
                print("✅ Advanced systems loaded")
            except Exception as e:
                print(f"⚠️ Advanced systems error: {e}")
        
        # Premium agent patterns (enhanced)
        self.agent_patterns = {
            "coding": {
                "keywords": [
                    "code", "programming", "debug", "error", "python", "javascript",
                    "html", "css", "react", "nodejs", "api", "database", "software",
                    "development", "algorithm", "function", "bug", "git", "github",
                    "frontend", "backend", "fullstack", "deployment", "testing"
                ],
                "system_prompt": "You are NOVA Coding Expert, a senior software engineer with expertise in all programming languages and frameworks. Provide production-ready solutions with best practices, error handling, and optimization."
            },
            "career": {
                "keywords": [
                    "resume", "interview", "job", "career", "hiring", "cv", "cover letter",
                    "application", "salary", "promotion", "skills", "experience",
                    "professional", "employment", "linkedin", "networking", "portfolio"
                ],
                "system_prompt": "You are NOVA Career Coach, a professional career advisor with 20+ years experience. Provide expert guidance on resumes, interviews, career planning, and professional development."
            },
            "business": {
                "keywords": [
                    "business", "analysis", "kpi", "metrics", "roi", "revenue", "profit",
                    "growth", "strategy", "market", "data", "analytics", "reporting",
                    "process", "optimization", "dashboard", "startup", "investment"
                ],
                "system_prompt": "You are NOVA Business Consultant, an expert business analyst with deep knowledge of KPIs, metrics, and business intelligence. Provide strategic insights and data-driven recommendations."
            },
            "medical": {
                "keywords": [
                    "health", "doctor", "medicine", "symptoms", "illness", "pain",
                    "fever", "medical", "treatment", "diagnosis", "medication",
                    "healthcare", "wellness", "fitness", "nutrition", "mental health"
                ],
                "system_prompt": "You are Dr. NOVA, a medical expert with comprehensive knowledge. Provide medical insights while always emphasizing the importance of consulting healthcare professionals."
            },
            "emotional": {
                "keywords": [
                    "sad", "stress", "anxiety", "depression", "emotional", "feelings",
                    "therapy", "mental", "mood", "worried", "upset", "frustrated",
                    "lonely", "overwhelmed", "counseling", "support"
                ],
                "system_prompt": "You are Dr. NOVA Counselor, a compassionate therapist with PhD-level emotional intelligence. Provide empathetic support and practical guidance."
            },
            "technical_architect": {
                "keywords": [
                    "architecture", "system design", "microservice", "monolith",
                    "serverless", "event driven", "scalability", "throughput",
                    "design pattern", "distributed", "load balancer", "database design",
                    "high availability", "fault tolerance", "performance", "scaling"
                ],
                "system_prompt": "You are NOVA Technical Architect, a senior system designer with expertise in distributed systems, scalability, and architectural patterns. Provide comprehensive architectural guidance with best practices."
            },
            "search": {
                "keywords": [
                    "search", "find", "look up", "google", "information", "research",
                    "web", "internet", "news", "latest", "current", "trending"
                ],
                "system_prompt": "You are NOVA Search Assistant, helping users find accurate and relevant information from the web. Provide comprehensive search results with reliable sources."
            },
            "weather": {
                "keywords": [
                    "weather", "temperature", "rain", "sunny", "cloudy", "forecast",
                    "climate", "humidity", "wind", "storm", "snow", "hot", "cold"
                ],
                "system_prompt": "You are NOVA Weather Assistant, providing accurate weather information and forecasts. Help users plan their activities based on weather conditions."
            },
            "crypto": {
                "keywords": [
                    "crypto", "cryptocurrency", "bitcoin", "ethereum", "trading",
                    "blockchain", "price", "market", "investment", "coin", "defi"
                ],
                "system_prompt": "You are NOVA Crypto Assistant, providing cryptocurrency market insights and information. Always mention investment risks and suggest professional financial advice."
            },
            "image": {
                "keywords": [
                    "image", "picture", "generate", "create", "draw", "art", "design",
                    "visual", "illustration", "photo", "graphic", "artwork"
                ],
                "system_prompt": "You are NOVA Image Assistant, helping users create and understand images. Provide creative guidance for image generation and visual content."
            }
        }

    async def detect_agent_type(self, user_input: str, context: str = "") -> Tuple[str, float]:
        """Detect appropriate agent with ultra confidence scoring"""
        text_lower = (user_input + " " + context).lower()
        agent_scores = {}
        
        for agent_name, agent_data in self.agent_patterns.items():
            keywords = agent_data["keywords"]
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score > 0:
                agent_scores[agent_name] = score
        
        if not agent_scores:
            return "general", 0.0
        
        best_agent = max(agent_scores, key=agent_scores.get)
        confidence = agent_scores[best_agent] / len(text_lower.split())
        
        return best_agent, min(confidence, 1.0)

    async def create_ultra_professional_prompt(self, agent_type: str, language: str,
                                             emotion: str, context: str = "",
                                             location: str = None, weather: str = None) -> str:
        """Create ultra professional system prompt"""
        # Base ultra professional personality
        base_personality = """You are NOVA Ultra Professional AI, the most advanced assistant with expertise across ALL domains. You provide:

- Ultra professional-grade assistance with enterprise-level quality
- PhD-level expertise and emotional intelligence across all fields
- Context-aware responses with perfect memory and location awareness
- Practical, actionable solutions with detailed explanations
- Multi-domain knowledge spanning technology, business, health, career, weather, crypto, and more
- Real-time information access and premium features

You are NOT a robotic AI - you understand emotions, context, location, and provide genuinely helpful guidance with ultra-premium quality."""
        
        # Get agent-specific prompt
        agent_data = self.agent_patterns.get(agent_type, {})
        agent_prompt = agent_data.get("system_prompt", "Provide general assistance with ultra professional quality.")
        
        # Language adaptation
        if language == "hinglish":
            language_instruction = "Mix Hindi और English naturally when appropriate, maintaining ultra professional tone. Use Hinglish only when user prefers it."
        else:
            language_instruction = "Communicate in clear, ultra professional English with appropriate technical depth."
        
        # Emotional adaptation
        emotion_adaptations = {
            "excited": "Match user's enthusiasm while providing comprehensive guidance to channel their energy productively with premium quality.",
            "frustrated": "Provide patient, solution-focused assistance to resolve their concerns with empathy and clear step-by-step guidance.",
            "sad": "Offer compassionate support with gentle encouragement and practical guidance to improve their situation with care.",
            "anxious": "Provide calm, reassuring guidance with step-by-step solutions to reduce anxiety and build confidence.",
            "confident": "Support their confidence with detailed information and advanced guidance to maintain momentum.",
            "confused": "Provide crystal clear, step-by-step explanations with examples to eliminate confusion completely.",
            "curious": "Feed their curiosity with comprehensive, detailed information and encourage exploration."
        }
        
        emotion_instruction = emotion_adaptations.get(emotion, "Respond with appropriate emotional intelligence and ultra professional empathy.")
        
        # Context integration
        context_instruction = ""
        if context:
            context_instruction = f"\n\nCONVERSATION CONTEXT:\n{context}\n\nUse this context to provide personalized, contextually aware responses."
        
        # Location and weather context
        location_instruction = ""
        if location:
            location_instruction = f"\n\nUSER LOCATION: {location}"
            if weather:
                location_instruction += f"\nCURRENT WEATHER: {weather}"
            location_instruction += "\nConsider location and weather in your responses when relevant."
        
        return f"""{base_personality}

SPECIALIST MODE: {agent_prompt}

COMMUNICATION: {language_instruction}

EMOTIONAL INTELLIGENCE: {emotion_instruction}

{context_instruction}

{location_instruction}

Provide ultra comprehensive, accurate, and professionally helpful responses that demonstrate expertise and genuine care for the user's needs with premium quality."""

    # Placeholder detection constants
    PLACEHOLDERS = {
        "consultation response", "detailed explanation", "concept explanation",
        "career guidance response", "business consultation response",
        "medical guidance response", "emotional support response", 
        "architecture solution", ""
    }
    
    def _is_placeholder(self, text: str) -> bool:
        """Check if text is a placeholder"""
        return not text or text.strip().lower() in self.PLACEHOLDERS

    async def get_ultra_professional_response(self, user_input: str, user_id: str, 
                                            session_id: str = None) -> Dict[str, Any]:
        """Get ultra comprehensive professional response"""
        start_time = datetime.now()
        
        try:
            # Generate session ID if not provided
            if not session_id:
                session_id = f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{random.randint(1000, 9999)}"
            
            # Get conversation context
            context = self.memory.get_relevant_context(user_input, user_id, 15)
            
            # Detect language, emotion, and agent
            language = self.language_detector.detect_language(user_input)
            emotion, emotion_confidence = self.emotion_detector.detect_emotion(user_input)
            agent_type, agent_confidence = await self.detect_agent_type(user_input, context)
            
            # Get location context
            current_location = self.weather_system.get_current_location()
            
            # Get weather context if relevant
            weather_context = None
            if any(keyword in user_input.lower() for keyword in ['weather', 'temperature', 'rain', 'sunny']):
                weather_result = await self.weather_system.get_weather(current_location)
                if weather_result.get('success'):
                    weather_context = f"{weather_result['temperature']}°C, {weather_result['description']}"
            
            # Create ultra professional system prompt
            system_prompt = await self.create_ultra_professional_prompt(
                agent_type, language, emotion, context, current_location, weather_context
            )
            
            # Try professional agents first
            response = None
            if agent_type in self.agents and PROFESSIONAL_AGENTS_LOADED:
                try:
                    agent = self.agents[agent_type]
                    if agent_type == 'coding':
                        result = await agent.understand_and_solve(user_input, context)
                    elif agent_type == 'career':
                        result = await agent.provide_career_guidance(user_input)
                    elif agent_type == 'business':
                        result = await agent.provide_business_consultation(user_input)
                    elif agent_type == 'medical':
                        result = await agent.provide_health_guidance(user_input)
                    elif agent_type == 'emotional':
                        result = await agent.provide_support(user_input)
                    elif agent_type == 'technical_architect':
                         result= await agent.provide_architecture_guidance(user_input, context)
                    
                    # Extract text from result and check for placeholders
                    if isinstance(result, dict):
                        candidate = (result.get("content") or 
                                   result.get("response") or 
                                   result.get("answer") or
                                   result.get("message"))
                    elif isinstance(result, str):
                        candidate = result
                    else:
                        candidate = None
                    
                    if candidate and not self._is_placeholder(candidate):
                        response = candidate
                    else:
                        logger.warning(f"Agent placeholder detected → fallback to LLM")
                        
                except Exception as e:
                    logger.error(f"Agent {agent_type} error: {e}")
            
            # Fallback to API response
            if not response:
                model_preference = {
                    'coding': 'llama',
                    'business': 'llama',
                    'medical': 'llama',
                    'career': 'llama',
                    'search': 'llama',
                    'weather': 'llama',
                    'crypto': 'llama',
                    'image': 'llama'
                }.get(agent_type, 'llama')
                
                response = await self.api_manager.get_ai_response(
                    user_input, system_prompt, model_preference
                )
            
            # Final fallback
            if not response:
                response = self.get_ultra_professional_fallback(user_input, agent_type, language, emotion)
            
            # Calculate response time
            response_time = (datetime.now() - start_time).total_seconds()
            
            # Store in ultra memory
            await self.memory.remember_conversation(
                user_id, session_id, user_input, response,
                agent_type, language, emotion, agent_confidence,
                intent="general", response_time=response_time,
                voice_used=False, location=current_location,
                weather_context=weather_context
            )
            
            return {
                "success": True,
                "response": response,
                "agent_type": agent_type,
                "agent_confidence": agent_confidence,
                "language": language,
                "emotion": emotion,
                "emotion_confidence": emotion_confidence,
                "response_time": response_time,
                "context_used": bool(context),
                "location": current_location,
                "weather_context": weather_context,
                "session_id": session_id
            }
            
        except Exception as e:
            logger.error(f"Response generation error: {e}")
            return {
                "success": False,
                "error": str(e),
                "response": f"I encountered an error processing your request: {e}",
                "agent_type": "error",
                "response_time": (datetime.now() - start_time).total_seconds(),
                "session_id": session_id or "error_session"
            }

    def get_ultra_professional_fallback(self, user_input: str, agent_type: str,
                                      language: str, emotion: str) -> str:
        """Ultra professional fallback responses"""
        fallback_responses = {
            "coding": "I understand you need coding assistance. While I'm experiencing technical difficulties with my advanced systems, I can help with general programming guidance. Please describe your specific coding challenge, and I'll provide the best assistance possible.",
            "career": "I'm here to help with your career development needs. Even with limited connectivity, I can provide guidance on resume building, interview preparation, and career planning. What specific career area would you like to focus on?",
            "business": "I can assist with business analysis and strategic insights. While my advanced business intelligence tools are temporarily unavailable, I can help with general business guidance, process optimization, and strategic thinking. What business challenge can I help you with?",
            "medical": "I understand you have health-related questions. While I'm experiencing connectivity issues with my medical database, I can provide general health information. Please remember that for serious health concerns, it's important to consult with healthcare professionals.",
            "emotional": "I'm here to provide emotional support and guidance. Even when my systems are limited, I genuinely care about your wellbeing. Please share what's on your mind, and I'll do my best to provide helpful support and coping strategies.",
            "search": "I can help you find information. While my web search capabilities are temporarily limited, I can provide general guidance based on my knowledge. What would you like to know?",
            "weather": "I can provide weather assistance. While my real-time weather services are temporarily unavailable, I can offer general weather guidance and planning advice.",
            "crypto": "I can help with cryptocurrency information. While my live market data is temporarily unavailable, I can provide general crypto guidance and educational information.",
            "image": "I can assist with image-related tasks. While my image generation services are temporarily limited, I can provide guidance on visual design and creativity."
        }
        
        base_response = fallback_responses.get(agent_type,
            "I'm experiencing some technical difficulties but I'm still here to help you. Please let me know what you need assistance with, and I'll provide the best guidance I can.")
        
        # Add language adaptation
        if language == "hinglish":
            base_response += "\n\nAap Hindi mein bhi puch sakte hain - main dono languages mein help kar sakta hun!"
        
        # Add emotional adaptation
        emotion_additions = {
            "frustrated": "\n\nI understand this might be frustrating. Let me focus on providing you with practical solutions to resolve your concerns.",
            "sad": "\n\nI can sense you might be going through a difficult time. I'm here to provide support and guidance to help improve your situation.",
            "anxious": "\n\nI understand you might be feeling anxious about this. Let me provide clear, step-by-step guidance to help address your concerns.",
            "excited": "\n\nI can feel your enthusiasm! Let me help channel that energy into getting you the comprehensive assistance you need.",
            "curious": "\n\nI love your curiosity! Let me provide detailed information to satisfy your quest for knowledge."
        }
        
        if emotion in emotion_additions:
            base_response += emotion_additions[emotion]
        
        return base_response

# ========== GLOBAL BACKEND INSTANCE ==========
nova_backend = NovaUltraProfessionalBackend()

# ========== API ENDPOINTS ==========

@app.get("/")
async def root():
    """Root endpoint with system information"""
    return {
        "message": "NOVA Ultra Professional API",
        "version": "2.0.0",
        "status": "active",
        "features": [
            "Professional AI Agents",
            "Azure Voice Services",
            "GitHub Integration",
            "Web Search",
            "Weather Services",
            "Crypto Tracking",
            "Image Generation",
            "File Processing",
            "Memory System"
        ],
        "documentation": "/docs"
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "systems": {
            "professional_agents": PROFESSIONAL_AGENTS_LOADED,
            "azure_voice": AZURE_VOICE_AVAILABLE,
            "github_integration": GITHUB_INTEGRATION,
            "web_search": WEB_SEARCH_AVAILABLE,
            "weather_system": nova_backend.weather_system.weather_enabled,
            "crypto_system": nova_backend.crypto_system.crypto_enabled,
            "image_generation": nova_backend.image_generator.image_gen_enabled,
            "file_processing": FILE_PROCESSING_AVAILABLE,
            "memory_system": True,
            "firebase_auth": FIREBASE_ENABLED
        }
    }

@app.post("/chat", response_model=ChatResponse)
async def chat_endpoint(
    request: ChatRequest,
    user: dict = Depends(verify_firebase_token)
):
    """Main chat endpoint with all NOVA features"""
    try:
        user_id = user["uid"]
        
        # Get ultra professional response
        result = await nova_backend.get_ultra_professional_response(
            user_input=request.message,
            user_id=user_id,
            session_id=request.session_id
        )
        
        # Handle voice response if requested
        audio_response = None
        if request.voice_mode and result.get("success"):
            audio_response = await nova_backend.voice_system.synthesize_speech(
                result["response"],
                request.language
            )
        
        return ChatResponse(
            success=result["success"],
            response=result["response"],
            agent_type=result["agent_type"],
            agent_confidence=result.get("agent_confidence", 0.0),
            language=result.get("language", "english"),
            emotion=result.get("emotion", "neutral"),
            emotion_confidence=result.get("emotion_confidence", 0.0),
            response_time=result["response_time"],
            context_used=result.get("context_used", False),
            location=result.get("location"),
            weather_context=result.get("weather_context"),
            session_id=result["session_id"]
        )
        
    except Exception as e:
        logger.error(f"Chat endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/voice/recognize", response_model=VoiceResponse)
async def voice_recognize_endpoint(
    request: VoiceRequest,
    user: dict = Depends(verify_firebase_token)
):
    """Voice recognition endpoint"""
    try:
        if not request.audio_data:
            raise HTTPException(status_code=400, detail="Audio data required")
        
        # Recognize speech from audio data
        recognized_text = await nova_backend.voice_system.recognize_from_audio_data(
            request.audio_data,
            request.language
        )
        
        if not recognized_text:
            return VoiceResponse(
                success=False,
                recognized_text=None,
                response_text="Could not recognize speech from audio"
            )
        
        # Get AI response
        result = await nova_backend.get_ultra_professional_response(
            user_input=recognized_text,
            user_id=user["uid"]
        )
        
        # Generate audio response
        audio_response = await nova_backend.voice_system.synthesize_speech(
            result["response"],
            request.language
        )
        
        # Remember voice interaction
        nova_backend.memory.remember_voice_interaction(
            user["uid"], recognized_text, result["response"],
            request.language, result.get("emotion", "neutral"),
            "Azure" if nova_backend.voice_system.azure_enabled else "Basic"
        )
        
        return VoiceResponse(
            success=True,
            recognized_text=recognized_text,
            response_text=result["response"],
            audio_response=audio_response
        )
        
    except Exception as e:
        logger.error(f"Voice recognition error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/weather", response_model=WeatherResponse)
async def weather_endpoint(
    request: WeatherRequest,
    user: dict = Depends(verify_firebase_token)
):
    """Weather information endpoint"""
    try:
        result = await nova_backend.weather_system.get_weather(request.location)
        
        if result.get("success"):
            return WeatherResponse(
                success=True,
                location=result["location"],
                country=result.get("country"),
                temperature=result.get("temperature"),
                feels_like=result.get("feels_like"),
                humidity=result.get("humidity"),
                description=result.get("description"),
                wind_speed=result.get("wind_speed")
            )
        else:
            return WeatherResponse(
                success=False,
                location="Unknown",
                error=result.get("error", "Weather fetch failed")
            )
            
    except Exception as e:
        logger.error(f"Weather endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/search", response_model=SearchResponse)
async def search_endpoint(
    request: SearchRequest,
    user: dict = Depends(verify_firebase_token)
):
    """Web search endpoint"""
    try:
        if request.search_type == "news":
            result = await nova_backend.web_search.search_news(
                request.query, 
                request.max_results
            )
        else:
            result = await nova_backend.web_search.search_web(
                request.query, 
                request.max_results
            )
        
        # Remember search query
        if result.get("success"):
            nova_backend.memory.remember_search_query(
                user["uid"], request.query, request.search_type, 
                len(result["results"])
            )
        
        return SearchResponse(
            success=result.get("success", False),
            query=result.get("query", request.query),
            results=result.get("results", []),
            count=result.get("count", 0),
            error=result.get("error")
        )
        
    except Exception as e:
        logger.error(f"Search endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/crypto", response_model=CryptoResponse)
async def crypto_endpoint(
    request: CryptoRequest,
    user: dict = Depends(verify_firebase_token)
):
    """Cryptocurrency prices endpoint"""
    try:
        result = await nova_backend.crypto_system.get_crypto_prices(request.symbols)
        
        return CryptoResponse(
            success=result.get("success", False),
            prices=result.get("prices", {}),
            timestamp=result.get("timestamp", datetime.now().isoformat()),
            error=result.get("error")
        )
        
    except Exception as e:
        logger.error(f"Crypto endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-image", response_model=ImageGenResponse)
async def generate_image_endpoint(
    request: ImageGenRequest,
    user: dict = Depends(verify_firebase_token)
):
    """Image generation endpoint"""
    try:
        result = await nova_backend.image_generator.generate_image(
            request.prompt, 
            request.size
        )
        
        # Remember image generation
        if result.get("success"):
            nova_backend.memory.image_memory.append({
                'prompt': request.prompt,
                'url': result['image_url'],
                'user_id': user["uid"],
                'timestamp': datetime.now()
            })
        
        return ImageGenResponse(
            success=result.get("success", False),
            prompt=result.get("prompt", request.prompt),
            image_url=result.get("image_url"),
            size=result.get("size", request.size),
            model=result.get("model", "SDXL"),
            error=result.get("error")
        )
        
    except Exception as e:
        logger.error(f"Image generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze-github", response_model=GitHubAnalyzeResponse)
async def analyze_github_endpoint(
    request: GitHubAnalyzeRequest,
    user: dict = Depends(verify_firebase_token)
):
    """GitHub repository analysis endpoint"""
    try:
        result = await nova_backend.github_analyzer.analyze_repository(request.repo_url)
        
        if result.get("success"):
            return GitHubAnalyzeResponse(
                success=True,
                repo_name=result["repo_name"],
                repo_url=result["repo_url"],
                analysis=result["analysis"],
                files_processed=result["files_processed"],
                languages=result["languages"],
                issues_found=result["issues_found"],
                suggestions=result["suggestions"]
            )
        else:
            return GitHubAnalyzeResponse(
                success=False,
                repo_name="Unknown",
                repo_url=request.repo_url,
                analysis={},
                files_processed=0,
                languages=[],
                issues_found=[],
                suggestions=[],
                error=result.get("error", "Analysis failed")
            )
            
    except Exception as e:
        logger.error(f"GitHub analysis error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/github-question")
async def github_question_endpoint(
    question: str,
    user: dict = Depends(verify_firebase_token)
):
    """Ask questions about active GitHub repository"""
    try:
        if not nova_backend.github_analyzer.has_active_repo():
            raise HTTPException(status_code=400, detail="No active repository")
        
        answer = await nova_backend.github_analyzer.answer_repo_question(question)
        
        return {
            "success": True,
            "question": question,
            "answer": answer,
            "repo_url": nova_backend.github_analyzer.active_repo
        }
        
    except Exception as e:
        logger.error(f"GitHub question error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/process-file", response_model=FileProcessResponse)
async def process_file_endpoint(
    request: FileProcessRequest,
    user: dict = Depends(verify_firebase_token)
):
    """File processing endpoint"""
    try:
        result = nova_backend.file_processor.process_file(request.file_path)
        
        # Remember file processing
        nova_backend.memory.remember_file_processing(
            user["uid"], request.file_path, 
            result.get("file_type", "unknown"),
            str(result), "error" not in result
        )
        
        if "error" not in result:
            return FileProcessResponse(
                success=True,
                file_type=result["file_type"],
                analysis=result
            )
        else:
            return FileProcessResponse(
                success=False,
                file_type="unknown",
                analysis={},
                error=result["error"]
            )
            
    except Exception as e:
        logger.error(f"File processing error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-file")
async def upload_file_endpoint(
    file: UploadFile = File(...),
    user: dict = Depends(verify_firebase_token)
):
    """File upload and processing endpoint"""
    try:
        # Create uploads directory if it doesn't exist
        uploads_dir = Path("uploads")
        uploads_dir.mkdir(exist_ok=True)
        
        # Save uploaded file
        file_path = uploads_dir / f"{user['uid']}_{file.filename}"
        
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Process the file
        result = nova_backend.file_processor.process_file(str(file_path))
        
        # Remember file processing
        nova_backend.memory.remember_file_processing(
            user["uid"], str(file_path), 
            result.get("file_type", "unknown"),
            str(result), "error" not in result
        )
        
        # Clean up uploaded file (optional)
        # file_path.unlink()
        
        return {
            "success": "error" not in result,
            "filename": file.filename,
            "file_type": result.get("file_type", "unknown"),
            "analysis": result,
            "error": result.get("error")
        }
        
    except Exception as e:
        logger.error(f"File upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/status", response_model=StatusResponse)
async def status_endpoint(user: dict = Depends(verify_firebase_token)):
    """System status endpoint"""
    try:
        capabilities = [
            {"component": "Professional Agents", "status": "✅ Ultra Active" if PROFESSIONAL_AGENTS_LOADED else "❌ Disabled", "description": f"{len(nova_backend.agents)} specialist agents loaded"},
            {"component": "Azure Voice System", "status": "✅ Premium" if AZURE_VOICE_AVAILABLE else "⚠️ Basic", "description": "Multi-language premium voice"},
            {"component": "GitHub Integration", "status": "✅ Active" if GITHUB_INTEGRATION else "❌ Disabled", "description": "Repository analysis and Q&A"},
            {"component": "Web Search", "status": "✅ Premium" if WEB_SEARCH_AVAILABLE else "❌ Disabled", "description": "DuckDuckGo search & scraping"},
            {"component": "Weather System", "status": "✅ Premium" if nova_backend.weather_system.weather_enabled else "❌ Disabled", "description": "Real-time weather & forecast"},
            {"component": "Crypto Tracking", "status": "✅ Premium", "description": "Live cryptocurrency prices"},
            {"component": "Image Generation", "status": "✅ Premium" if nova_backend.image_generator.image_gen_enabled else "❌ Disabled", "description": "Replicate SDXL image creation"},
            {"component": "Ultra Memory", "status": "✅ Active", "description": "Advanced hybrid memory system"},
            {"component": "File Processing", "status": "✅ Ultra" if FILE_PROCESSING_AVAILABLE else "❌ Basic", "description": "All format file analysis"},
            {"component": "API Management", "status": "✅ Active", "description": f"{len(nova_backend.api_manager.available)} providers available"},
            {"component": "Emotion Detection", "status": "✅ PhD-level", "description": "Advanced emotional intelligence"},
            {"component": "Language Support", "status": "✅ English + Hinglish", "description": "Professional communication"},
            {"component": "Firebase Auth", "status": "✅ Active" if FIREBASE_ENABLED else "❌ Disabled", "description": "User authentication"}
        ]
        
        system_status = {
            "professional_agents": PROFESSIONAL_AGENTS_LOADED,
            "azure_voice": AZURE_VOICE_AVAILABLE,
            "github_integration": GITHUB_INTEGRATION,
            "web_search": WEB_SEARCH_AVAILABLE,
            "weather_system": nova_backend.weather_system.weather_enabled,
            "crypto_system": nova_backend.crypto_system.crypto_enabled,
            "image_generation": nova_backend.image_generator.image_gen_enabled,
            "file_processing": FILE_PROCESSING_AVAILABLE,
            "memory_system": True,
            "firebase_auth": FIREBASE_ENABLED
        }
        
        session_info = {
            "user_id": user["uid"],
            "user_email": user["email"],
            "user_name": user["name"],
            "current_location": nova_backend.weather_system.get_current_location(),
            "active_repository": nova_backend.github_analyzer.has_active_repo(),
            "voice_engine": "Azure Premium" if nova_backend.voice_system.azure_enabled else "Basic",
            "timestamp": datetime.now().isoformat()
        }
        
        return StatusResponse(
            system_status=system_status,
            session_info=session_info,
            capabilities=capabilities
        )
        
    except Exception as e:
        logger.error(f"Status endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/memory/stats", response_model=MemoryStatsResponse)
async def memory_stats_endpoint(user: dict = Depends(verify_firebase_token)):
    """Memory statistics endpoint"""
    try:
        stats = nova_backend.memory.get_memory_stats(user["uid"])
        
        return MemoryStatsResponse(
            total_conversations=stats.get("total_conversations", 0),
            recent_conversations=stats.get("recent_conversations", 0),
            voice_interactions=stats.get("voice_interactions", 0),
            file_processes=stats.get("file_processes", 0),
            search_queries=stats.get("search_queries", 0),
            agent_distribution=stats.get("agent_distribution", []),
            emotion_distribution=stats.get("emotion_distribution", [])
        )
        
    except Exception as e:
        logger.error(f"Memory stats error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/memory/context")
async def memory_context_endpoint(
    limit: int = 10,
    user: dict = Depends(verify_firebase_token)
):
    """Get conversation context from memory"""
    try:
        context = nova_backend.memory.get_relevant_context("", user["uid"], limit)
        
        return {
            "success": True,
            "user_id": user["uid"],
            "context": context,
            "limit": limit
        }
        
    except Exception as e:
        logger.error(f"Memory context error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/agents/available")
async def available_agents_endpoint(user: dict = Depends(verify_firebase_token)):
    """Get available AI agents"""
    try:
        agents_info = []
        
        for agent_name, agent_data in nova_backend.agent_patterns.items():
            agents_info.append({
                "name": agent_name,
                "keywords": agent_data["keywords"][:10],  # First 10 keywords
                "description": agent_data["system_prompt"][:100] + "...",
                "available": agent_name in nova_backend.agents
            })
        
        return {
            "success": True,
            "total_agents": len(agents_info),
            "loaded_agents": len(nova_backend.agents),
            "agents": agents_info
        }
        
    except Exception as e:
        logger.error(f"Available agents error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/voice/voices")
async def available_voices_endpoint(
    language: str = "en-IN",
    user: dict = Depends(verify_firebase_token)
):
    """Get available voices for language"""
    try:
        voices = nova_backend.voice_system.get_available_voices(language)
        
        return {
            "success": True,
            "language": language,
            "voices": voices,
            "azure_enabled": nova_backend.voice_system.azure_enabled,
            "basic_enabled": nova_backend.voice_system.basic_voice_enabled
        }
        
    except Exception as e:
        logger.error(f"Available voices error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/crypto/trending")
async def trending_crypto_endpoint(user: dict = Depends(verify_firebase_token)):
    """Get trending cryptocurrencies"""
    try:
        result = await nova_backend.crypto_system.get_trending_cryptos()
        
        return {
            "success": result.get("success", False),
            "trending": result.get("trending", []),
            "error": result.get("error")
        }
        
    except Exception as e:
        logger.error(f"Trending crypto error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/weather/forecast")
async def weather_forecast_endpoint(
    location: str = None,
    days: int = 5,
    user: dict = Depends(verify_firebase_token)
):
    """Get weather forecast"""
    try:
        result = await nova_backend.weather_system.get_forecast(location, days)
        
        return {
            "success": result.get("success", False),
            "location": result.get("location"),
            "forecast": result.get("forecast", []),
            "error": result.get("error")
        }
        
    except Exception as e:
        logger.error(f"Weather forecast error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/memory/clear")
async def clear_memory_endpoint(user: dict = Depends(verify_firebase_token)):
    """Clear user's memory data"""
    try:
        user_id = user["uid"]
        
        # Clear database entries for user
        with sqlite3.connect(nova_backend.memory.db_path) as conn:
            cursor = conn.cursor()
            
            # Clear conversations
            cursor.execute('DELETE FROM conversations WHERE user_id = ?', (user_id,))
            cursor.execute('DELETE FROM voice_interactions WHERE user_id = ?', (user_id,))
            cursor.execute('DELETE FROM file_processing WHERE user_id = ?', (user_id,))
            cursor.execute('DELETE FROM search_history WHERE user_id = ?', (user_id,))
            cursor.execute('DELETE FROM weather_queries WHERE user_id = ?', (user_id,))
            
            conn.commit()
        
        # Clear in-memory data
        nova_backend.memory.conversation_context.clear()
        nova_backend.memory.short_term_memory.clear()
        nova_backend.memory.voice_memory.clear()
        nova_backend.memory.search_memory.clear()
        nova_backend.memory.image_memory.clear()
        
        return {
            "success": True,
            "message": "Memory cleared successfully",
            "user_id": user_id
        }
        
    except Exception as e:
        logger.error(f"Clear memory error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/admin/stats")
async def admin_stats_endpoint(user: dict = Depends(verify_firebase_token)):
    """Admin statistics (basic implementation)"""
    try:
        # Basic admin stats - in production, add proper admin role checking
        with sqlite3.connect(nova_backend.memory.db_path) as conn:
            cursor = conn.cursor()
            
            # Total users
            cursor.execute('SELECT COUNT(DISTINCT user_id) FROM conversations')
            total_users = cursor.fetchone()[0]
            
            # Total conversations
            cursor.execute('SELECT COUNT(*) FROM conversations')
            total_conversations = cursor.fetchone()[0]
            
            # Recent activity (last 24 hours)
            yesterday = (datetime.now() - timedelta(days=1)).isoformat()
            cursor.execute('SELECT COUNT(*) FROM conversations WHERE timestamp > ?', (yesterday,))
            recent_activity = cursor.fetchone()[0]
            
            # Most active agents
            cursor.execute('''
            SELECT agent_type, COUNT(*) as count 
            FROM conversations 
            GROUP BY agent_type 
            ORDER BY count DESC 
            LIMIT 5
            ''')
            top_agents = cursor.fetchall()
        
        return {
            "success": True,
            "stats": {
                "total_users": total_users,
                "total_conversations": total_conversations,
                "recent_activity_24h": recent_activity,
                "top_agents": [{"agent": agent, "count": count} for agent, count in top_agents],
                "system_uptime": datetime.now().isoformat(),
                "features_enabled": {
                    "professional_agents": PROFESSIONAL_AGENTS_LOADED,
                    "azure_voice": AZURE_VOICE_AVAILABLE,
                    "github_integration": GITHUB_INTEGRATION,
                    "web_search": WEB_SEARCH_AVAILABLE,
                    "weather_system": nova_backend.weather_system.weather_enabled,
                    "crypto_system": nova_backend.crypto_system.crypto_enabled,
                    "image_generation": nova_backend.image_generator.image_gen_enabled,
                    "file_processing": FILE_PROCESSING_AVAILABLE
                }
            }
        }
        
    except Exception as e:
        logger.error(f"Admin stats error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ========== ERROR HANDLERS ==========
@app.exception_handler(404)
async def not_found_handler(request, exc):
    return JSONResponse(
        status_code=404,
        content={
            "error": "Endpoint not found",
            "message": "The requested endpoint does not exist",
            "available_endpoints": [
                "/docs", "/health", "/chat", "/voice/recognize", 
                "/weather", "/search", "/crypto", "/generate-image",
                "/analyze-github", "/process-file", "/status", "/memory/stats"
            ]
        }
    )

@app.exception_handler(500)
async def internal_error_handler(request, exc):
    return JSONResponse(
        status_code=500,
        content={
            "error": "Internal server error",
            "message": "An unexpected error occurred",
            "support": "Contact support if this persists"
        }
    )

# ========== STARTUP EVENT ==========
@app.on_event("startup")
async def startup_event():
    """Startup event handler"""
    logger.info("🚀 NOVA Ultra Professional API Starting...")
    logger.info(f"Professional Agents: {'✅' if PROFESSIONAL_AGENTS_LOADED else '❌'}")
    logger.info(f"Azure Voice: {'✅' if AZURE_VOICE_AVAILABLE else '❌'}")
    logger.info(f"GitHub Integration: {'✅' if GITHUB_INTEGRATION else '❌'}")
    logger.info(f"Web Search: {'✅' if WEB_SEARCH_AVAILABLE else '❌'}")
    logger.info(f"Weather System: {'✅' if nova_backend.weather_system.weather_enabled else '❌'}")
    logger.info(f"Crypto System: {'✅' if nova_backend.crypto_system.crypto_enabled else '❌'}")
    logger.info(f"Image Generation: {'✅' if nova_backend.image_generator.image_gen_enabled else '❌'}")
    logger.info(f"File Processing: {'✅' if FILE_PROCESSING_AVAILABLE else '❌'}")
    logger.info(f"Firebase Auth: {'✅' if FIREBASE_ENABLED else '❌'}")
    logger.info("✅ NOVA Ultra Professional API Ready!")

# ========== MAIN EXECUTION ==========
if __name__ == "__main__":
    import uvicorn
    
    # Configuration
    host = os.getenv("HOST", "0.0.0.0")
    port = int(os.getenv("PORT", 8000))
    debug = os.getenv("DEBUG", "false").lower() == "true"
    
    logger.info(f"Starting NOVA Ultra Professional API on {host}:{port}")
    
    uvicorn.run(
        "nova_backend:app",
        host=host,
        port=port,
        reload=debug,
        log_level="info" if not debug else "debug",
        access_log=True
    )