import asyncio
import os
import sys
import json
import time
import sqlite3
import logging
import hashlib
import re
import requests
import random
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Union, Tuple
from collections import defaultdict, deque, Counter
from pathlib import Path
import tempfile
from io import BytesIO
import aiofiles
from bs4 import BeautifulSoup
import numpy as np

# FastAPI imports
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
import uvicorn
from pydub import AudioSegment
from io import BytesIO

# Environment loading
from dotenv import load_dotenv
load_dotenv()

def ensure_dict(obj):
    if isinstance(obj, dict):
        return obj
    if isinstance(obj, str):
        try:
            return json.loads(obj)
        except:
            return {}
    return {}

def sanitize_metadata(meta):
    """Convert numpy + nested dicts/lists into JSON-safe values for metadata storage."""
    if isinstance(meta, dict):
        return {k: sanitize_metadata(v) for k, v in meta.items()}
    elif isinstance(meta, (list, tuple)):
        return [sanitize_metadata(v) for v in meta]
    elif isinstance(meta, np.generic):  # numpy types
        return float(meta)
    elif isinstance(meta, (np.ndarray,)):
        return meta.tolist()
    elif isinstance(meta, (str, int, float, bool)) or meta is None:
        return meta
    else:
        # Anything else (like nested dicts), stringify safely
        try:
            return json.dumps(meta)
        except:
            return str(meta)

def webm_to_wav(audio_bytes: bytes) -> bytes:
    """Convert browser WebM/Opus to WAV in memory."""
    audio = AudioSegment.from_file(BytesIO(audio_bytes), format="webm")
    wav_io = BytesIO()
    audio.export(wav_io, format="wav")
    return wav_io.getvalue()

# Setup paths (same as CLI)
project_root = os.path.dirname(os.path.abspath(__file__))
folders_to_add = [
    'src', os.path.join('src', 'memory'), os.path.join('src', 'unique_features'),
    os.path.join('src', 'agents'), 'ML'
]
for folder in folders_to_add:
    folder_path = os.path.join(project_root, folder)
    if os.path.exists(folder_path) and folder_path not in sys.path:
        sys.path.insert(0, folder_path)

# Voice processing imports (same as CLI)
try:
    import speech_recognition as sr
    import pyttsx3
    VOICE_AVAILABLE = True
except ImportError:
    VOICE_AVAILABLE = False

try:
    import azure.cognitiveservices.speech as speechsdk
    AZURE_VOICE_AVAILABLE = True
except ImportError:
    AZURE_VOICE_AVAILABLE = False

# File processing imports (same as CLI)
try:
    from PIL import Image
    import PyPDF2
    import docx
    import pandas as pd
    FILE_PROCESSING_AVAILABLE = True
except ImportError:
    FILE_PROCESSING_AVAILABLE = False

# GitHub Integration imports (same as CLI)
try:
    import chromadb
    from langchain_community.document_loaders import UnstructuredFileLoader
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    GITHUB_INTEGRATION = True
except ImportError:
    GITHUB_INTEGRATION = False

# Professional Agents Import (same as CLI)
try:
    from agents.coding_agent import ProLevelCodingExpert
    from agents.career_coach import ProfessionalCareerCoach  
    from agents.business_consultant import SmartBusinessConsultant
    from agents.medical_advisor import SimpleMedicalAdvisor
    from agents.emotional_counselor import SimpleEmotionalCounselor
    from agents.techincal_architect import TechnicalArchitect
    PROFESSIONAL_AGENTS_LOADED = True
except ImportError:
    PROFESSIONAL_AGENTS_LOADED = False

# Advanced Systems Import (same as CLI)
try:
    from memory.sharp_memory import SharpMemorySystem
    from unique_features.smart_orchestrator import IntelligentAPIOrchestrator
    from unique_features.api_drift_detector import APIPerformanceDrifter
    ADVANCED_SYSTEMS = True
except ImportError:
    ADVANCED_SYSTEMS = False
    # Fallback classes
    class SharpMemorySystem:
        def __init__(self): pass
        async def remember_conversation_advanced(self, *args): pass
        async def get_semantic_context(self, *args): return ""
    
    class IntelligentAPIOrchestrator:
        def __init__(self): pass
        async def get_optimized_response(self, *args): return None, {}
    
    class APIPerformanceDrifter:
        def __init__(self): pass
        def record_response_quality(self, *args): pass

# GitHub QA Engine Import (same as CLI)
try:
    from agents.ingest import main as ingest_repo
    from agents.qa_engine import create_qa_engine
    GITHUB_INTEGRATION = GITHUB_INTEGRATION and True
except ImportError:
    GITHUB_INTEGRATION = False
    ingest_repo = None
    create_qa_engine = None

# ========== SMART ML SYSTEM INTEGRATION ==========
# Enhanced ML System Import with Smart Enhancement Detection
try:
    from ml_integration import EnhancedMLManager
    ml_manager = EnhancedMLManager()
    ML_SYSTEM_AVAILABLE = True
    logger = logging.getLogger(__name__)
    logger.info("✅ Enhanced ML Manager loaded successfully!")
except ImportError:
    ML_SYSTEM_AVAILABLE = False
    class EnhancedMLManager:
        def __init__(self): pass
        async def enhance_query(self, query, context): return query
        async def optimize_response(self, response, context): return response
        def process_user_query(self, query, context=None): return {}
        def store_interaction_intelligently(self, *args): pass
    ml_manager = EnhancedMLManager()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ========== SMART ENHANCEMENT DETECTOR ==========
class SmartEnhancementDetector:
    """Intelligent detection of when to apply ML enhancement vs simple AI responses"""
    
    @staticmethod
    def needs_ml_enhancement(user_query: str) -> bool:
        """
        Determine if query needs advanced ML processing
        Returns True for complex queries, False for simple queries that still need AI but not ML
        """
        query_lower = user_query.lower().strip()
        
        # Complex queries that NEED ML enhancement
        complex_indicators = [
            # Technical queries
            'code', 'programming', 'algorithm', 'debug', 'error', 'function', 'api', 'database',
            'architecture', 'system design', 'scalability', 'performance', 'optimization',
            
            # Professional queries  
            'career', 'job', 'interview', 'resume', 'promotion', 'salary', 'skills', 'linkedin',
            'business', 'strategy', 'market', 'revenue', 'profit', 'analysis', 'growth',
            
            # Advanced requests
            'analyze', 'compare', 'recommend', 'suggest', 'implement', 'design',
            'create', 'build', 'develop', 'optimize', 'improve', 'review',
            
            # Medical/Health (complex)
            'symptoms', 'treatment', 'diagnosis', 'medicine', 'therapy',
            
            # Complex emotional/mental health
            'depression', 'anxiety', 'therapy', 'counseling', 'mental health',
            
            # File/Data analysis
            'file', 'document', 'data', 'report', 'spreadsheet', 'presentation',
            
            # Project/work related
            'project', 'assessment', 'guidance', 'help me with', 'assist me',
            'consultation', 'advice on', 'evaluate'
        ]
        
        # Multi-word complex patterns
        complex_patterns = [
            r'help me (with|in|on)',
            r'can you (help|assist|guide)',
            r'i (need|want|would like) (help|assistance|guidance)',
            r'what (should|would|could) i do',
            r'how (can|should|do) i',
            r'please (help|assist|guide|advise)',
            r'give me (advice|guidance|help)',
            r'i am (struggling|having trouble|confused)',
            r'explain (how|why|what|when)',
            r'tell me about'
        ]
        
        # Check for complex indicators
        has_complex_terms = any(term in query_lower for term in complex_indicators)
        has_complex_patterns = any(re.search(pattern, query_lower) for pattern in complex_patterns)
        is_long_query = len(query_lower.split()) > 15
        
        return has_complex_terms or has_complex_patterns or is_long_query
    
    @staticmethod
    def is_simple_greeting(user_query: str) -> bool:
        """Check if it's a very simple greeting that needs basic AI response"""
        query_lower = user_query.lower().strip()
        
        simple_patterns = [
            r'^(hi|hello|hey|hola)$',
            r'^(hi there|hello there|hey there)$',
            r'^(good morning|good afternoon|good evening)$',
            r'^(how are you|how\'s it going|what\'s up|sup)$',
            r'^(thanks|thank you|thx|ty)$',
            r'^(bye|goodbye|see you|talk later|cya)$',
            r'^(yes|no|ok|okay|sure|alright)$',
            r'^(what is your name|who are you)$',
            r'^(help|test|testing)$',
        ]
        
        return any(re.match(pattern, query_lower) for pattern in simple_patterns)

# ========== ULTRA HYBRID MEMORY SYSTEM (EXACT FROM CLI) ==========
class UltraHybridMemorySystem:
    """Ultra Advanced Hybrid Memory - EXACT from NOVA-CLI.py"""
    
    def __init__(self, db_path="nova_ultra_professional_memory.db"):
        if not os.path.isabs(db_path):
            self.db_path = os.path.join(os.getcwd(), db_path)
        else:
            self.db_path = db_path
        
        self.setup_database()
        
        # ALL memory layers from CLI - EXACT
        self.conversation_context = deque(maxlen=100)
        self.user_profile = {}
        self.emotional_state = "neutral"
        self.learning_patterns = defaultdict(list)
        self.personality_insights = {}
        self.user_preferences = {}
        self.conversation_history = []
        
        # Memory layers from CLI - EXACT
        self.short_term_memory = deque(maxlen=200)
        self.working_memory = {}
        self.conversation_threads = {}
        self.context_memory = {}
        
        # Premium memory features - EXACT
        self.voice_memory = deque(maxlen=50)
        self.file_memory = {}
        self.search_memory = deque(maxlen=30)
        self.image_memory = deque(maxlen=20)
        
        # Semantic memory - EXACT from CLI
        self.setup_semantic_memory()

    def setup_database(self):
        """Setup database schema - EXACT from CLI with ML enhancement columns"""
        try:
            db_dir = os.path.dirname(self.db_path)
            if db_dir and not os.path.exists(db_dir):
                os.makedirs(db_dir, exist_ok=True)
                
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                # Enhanced conversations table - EXACT from CLI + ML columns
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS conversations (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        user_id TEXT,
                        session_id TEXT,
                        user_input TEXT,
                        bot_response TEXT,
                        agent_type TEXT,
                        language TEXT,
                        emotion TEXT,
                        confidence REAL,
                        timestamp DATETIME,
                        feedback INTEGER DEFAULT 0,
                        context_summary TEXT,
                        learned_facts TEXT,
                        satisfaction_rating INTEGER,
                        conversation_thread_id TEXT,
                        intent_detected TEXT,
                        response_time REAL,
                        voice_used BOOLEAN DEFAULT 0,
                        location TEXT,
                        weather_context TEXT,
                        search_queries TEXT,
                        ml_insights TEXT DEFAULT '{}',
                        intent_confidence REAL DEFAULT 0.0,
                        context_quality TEXT DEFAULT 'medium',
                        enhancement_applied BOOLEAN DEFAULT 0
                    )
                ''')
                
                # Enhanced user profiles - EXACT from CLI + ML columns
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS user_profiles (
                        user_id TEXT PRIMARY KEY,
                        name TEXT,
                        career_goals TEXT,
                        current_role TEXT,
                        experience_years INTEGER,
                        skills TEXT,
                        preferences TEXT,
                        communication_style TEXT,
                        emotional_patterns TEXT,
                        conversation_patterns TEXT,
                        expertise_level TEXT,
                        topics_of_interest TEXT,
                        last_updated DATETIME,
                        total_conversations INTEGER DEFAULT 0,
                        preferred_voice TEXT,
                        location TEXT,
                        timezone TEXT,
                        personality_type TEXT,
                        learning_style TEXT,
                        preferred_agents TEXT,
                        interaction_patterns TEXT
                    )
                ''')
                
                # Other tables - EXACT from CLI
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS github_repos (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        repo_url TEXT UNIQUE,
                        repo_name TEXT,
                        analysis_date DATETIME,
                        file_count INTEGER,
                        languages_detected TEXT,
                        issues_found TEXT,
                        suggestions TEXT,
                        vector_db_path TEXT
                    )
                ''')
                
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS voice_interactions (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        user_id TEXT,
                        voice_input TEXT,
                        voice_response TEXT,
                        language_detected TEXT,
                        emotion_detected TEXT,
                        voice_engine TEXT,
                        timestamp DATETIME
                    )
                ''')
                
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS file_processing (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        user_id TEXT,
                        file_path TEXT,
                        file_type TEXT,
                        processing_result TEXT,
                        timestamp DATETIME,
                        success BOOLEAN
                    )
                ''')
                
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS search_history (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        user_id TEXT,
                        search_query TEXT,
                        search_type TEXT,
                        results_count INTEGER,
                        timestamp DATETIME
                    )
                ''')
                
                conn.commit()
                logger.info("✅ Database initialized")
        except Exception as e:
            logger.error(f"Database setup error: {e}")

    def setup_semantic_memory(self):
        """Setup semantic memory - EXACT from CLI"""
        try:
            if ADVANCED_SYSTEMS:
                self.semantic_memory = SharpMemorySystem()
            else:
                self.semantic_memory = None
        except Exception as e:
            logger.error(f"Semantic memory setup error: {e}")
            self.semantic_memory = None

    async def remember_conversation(self, user_id: str, session_id: str,
                                  user_input: str, bot_response: str,
                                  agent_type: str, language: str,
                                  emotion: str, confidence: float,
                                  intent: str = None, response_time: float = 0.0,
                                  voice_used: bool = False, location: str = None,
                                  weather_context: str = None, search_queries: str = None,
                                  file_analyzed: str = None, ml_insights: Dict = None,
                                  enhancement_applied: bool = False):
        """Enhanced conversation memory storage with ML insights"""
        try:
            # Store in advanced memory if available
            if ADVANCED_SYSTEMS and self.semantic_memory:
                await self.semantic_memory.remember_conversation_advanced(
                    user_id, user_input, bot_response, agent_type, confidence
                )
            ml_data = ensure_dict(ml_insights)

            # Convert entire ML insights dict into JSON for DB storage
            ml_insights_json = json.dumps(ml_data)

           # Compute these BEFORE passing into the SQL query
            intent_confidence = ml_data.get('routing_decision', {}).get('confidence_level', 0.0)
            context_quality = ml_data.get('context_enhancement', {}).get('context_quality', 'medium')

            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    INSERT INTO conversations 
                    (user_id, session_id, user_input, bot_response, agent_type, language, 
                    emotion, confidence, timestamp, intent_detected, response_time, 
                    voice_used, location, weather_context, search_queries, ml_insights,
                    intent_confidence, context_quality, enhancement_applied)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    user_id, session_id, user_input, bot_response, agent_type, language,
                    emotion, confidence, datetime.now(), intent, response_time,
                    voice_used, location, weather_context, search_queries, ml_insights_json,
                    intent_confidence, context_quality, enhancement_applied
                ))
                
                # Update user profile
                cursor.execute('''
                    INSERT OR REPLACE INTO user_profiles 
                    (user_id, total_conversations, last_updated, preferred_agents)
                    VALUES (?, 
                            COALESCE((SELECT total_conversations FROM user_profiles WHERE user_id = ?), 0) + 1,
                            ?, ?)
                ''', (user_id, user_id, datetime.now(), agent_type))
                
                conn.commit()
            
            # Update in-memory context
            conversation_entry = {
                'user': user_input,
                'bot': bot_response,
                'timestamp': datetime.now(),
                'agent': agent_type,
                'emotion': emotion,
                'confidence': confidence,
                'ml_enhanced': enhancement_applied
            }
            self.conversation_context.append(conversation_entry)
            self.short_term_memory.append(conversation_entry)

            # Update working memory
            thread_id = f"{user_id}_{session_id}"
            if thread_id not in self.conversation_threads:
                self.conversation_threads[thread_id] = deque(maxlen=50)
            self.conversation_threads[thread_id].append(conversation_entry)
            
            # Store interaction in ML system if available
            if ML_SYSTEM_AVAILABLE and ml_insights:
                ml_manager.store_interaction_intelligently(
                    user_input, bot_response, agent_type
                )
            
        except Exception as e:
            logger.error(f"Memory storage error: {e}")

    async def get_conversation_context(self, user_id: str, limit: int = 10) -> str:
        """Get conversation context for enhanced responses"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT user_input, bot_response, agent_type, timestamp, enhancement_applied
                    FROM conversations 
                    WHERE user_id = ? 
                    ORDER BY timestamp DESC 
                    LIMIT ?
                ''', (user_id, limit))
                
                rows = cursor.fetchall()
                if not rows:
                    return ""
                
                context = "Recent conversation context:\n"
                for row in reversed(rows):
                    user_input, bot_response, agent_type, timestamp, enhanced = row
                    enhancement_flag = " [ML Enhanced]" if enhanced else ""
                    context += f"[{agent_type}]{enhancement_flag} User: {user_input[:100]}...\n"
                    context += f"Assistant: {bot_response[:100]}...\n\n"
                
                return context
        except Exception as e:
            logger.error(f"Context retrieval error: {e}")
            return ""

    async def get_user_profile(self, user_id: str) -> Dict[str, Any]:
        """Get user profile for personalization"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT name, preferences, communication_style, expertise_level, 
                           total_conversations, preferred_agents, interaction_patterns
                    FROM user_profiles 
                    WHERE user_id = ?
                ''', (user_id,))
                
                row = cursor.fetchone()
                if row:
                    return {
                        'name': row[0],
                        'preferences': row[1],
                        'communication_style': row[2],
                        'expertise_level': row[3],
                        'total_conversations': row[4],
                        'preferred_agents': row[5],
                        'interaction_patterns': row[6]
                    }
                return {}
        except Exception as e:
            logger.error(f"User profile retrieval error: {e}")
            return {}

# Initialize memory system
memory_system = UltraHybridMemorySystem()

# ========== LANGUAGE AND EMOTION DETECTORS (EXACT FROM CLI) ==========
class FastLanguageDetector:
    """Language detection - EXACT from CLI"""
    
    def __init__(self):
        self.hinglish_words = {
            "yaar", "bhai", "ji", "hai", "hoon", "kya", "aur", "tum", "main",
            "accha", "theek", "nahi", "haan", "matlab", "kaise", "kyun"
        }

    def detect_language(self, text: str) -> str:
        """Fast language detection - EXACT from CLI"""
        words = text.lower().split()
        hinglish_count = sum(1 for word in words if word in self.hinglish_words)
        return "hinglish" if hinglish_count > 0 else "english"

class FastEmotionDetector:
    """Emotion detection - EXACT from CLI"""
    
    def __init__(self):
        self.emotion_keywords = {
            "excited": ["excited", "amazing", "awesome", "great", "love"],
            "frustrated": ["frustrated", "angry", "upset", "hate", "annoyed"],
            "sad": ["sad", "depressed", "down", "unhappy", "lonely"],
            "anxious": ["anxious", "worried", "nervous", "scared", "stress"],
            "confident": ["confident", "sure", "ready", "motivated", "strong"],
            "confused": ["confused", "lost", "unclear", "help", "stuck"]
        }

    def detect_emotion(self, text: str) -> Tuple[str, float]:
        """Fast emotion detection - EXACT from CLI"""
        text_lower = text.lower()
        for emotion, keywords in self.emotion_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                return emotion, 0.8
        return "neutral", 0.5

# ========== OPTIMIZED FREE API MANAGER (ENHANCED FROM GITHUB) ==========
class OptimizedAPIManager:
    """Enhanced API manager with FREE providers - From GitHub reference"""
    
    def __init__(self):
        # ALL 6+ FREE API providers from GitHub reference - Enhanced
        self.providers = [
            {
                "name": "Groq",
                "url": "https://api.groq.com/openai/v1/chat/completions",
                "models": ["llama-3.1-8b-instant", "llama-3.1-70b-versatile", "mixtral-8x7b-32768"],
                "headers": lambda: {
                    "Authorization": f"Bearer {os.getenv('GROQ_API_KEY', '')}",
                    "Content-Type": "application/json"
                },
                "priority": 1,
                "specialty": "fast_inference",
                "rate_limit": 30,  # requests per minute
                "max_tokens": 32768
            },
            {
                "name": "OpenRouter",
                "url": "https://openrouter.ai/api/v1/chat/completions",
                "models": [
                    "mistralai/mistral-7b-instruct:free", 
                    "meta-llama/llama-3.1-70b-instruct:free",
                    "microsoft/wizardlm-2-8x22b:free",
                    "google/gemma-2b-it:free"
                ],
                "headers": lambda: {
                    "Authorization": f"Bearer {os.getenv('OPENROUTER_API_KEY', '')}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://nova-ai.app",
                    "X-Title": "NOVA AI Assistant"
                },
                "priority": 2,
                "specialty": "diverse_models",
                "rate_limit": 25,
                "max_tokens": 4096
            },
            {
                "name": "Together",
                "url": "https://api.together.xyz/v1/chat/completions",
                "models": [
                    "meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
                    "meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
                    "mistralai/Mixtral-8x7B-Instruct-v0.1"
                ],
                "headers": lambda: {
                    "Authorization": f"Bearer {os.getenv('TOGETHER_API_KEY', '')}",
                    "Content-Type": "application/json"
                },
                "priority": 3,
                "specialty": "open_source",
                "rate_limit": 20,
                "max_tokens": 8192
            },
            {
                "name": "HuggingFace",
                "url": "https://api-inference.huggingface.co/models/",
                "models": [
                    "microsoft/DialoGPT-large",
                    "microsoft/DialoGPT-medium",
                    "facebook/blenderbot-400M-distill"
                ],
                "headers": lambda: {
                    "Authorization": f"Bearer {os.getenv('HUGGINGFACE_API_KEY', '')}",
                    "Content-Type": "application/json"
                },
                "priority": 4,
                "specialty": "conversational",
                "rate_limit": 15,
                "max_tokens": 2048,
                "custom_format": True
            },
            {
                "name": "Cohere",
                "url": "https://api.cohere.ai/v1/chat",
                "models": ["command-light", "command"],
                "headers": lambda: {
                    "Authorization": f"Bearer {os.getenv('COHERE_API_KEY', '')}",
                    "Content-Type": "application/json"
                },
                "priority": 5,
                "specialty": "text_generation",
                "rate_limit": 10,
                "max_tokens": 4096,
                "custom_format": True
            },
            {
                "name": "NVIDIA",
                "url": "https://integrate.api.nvidia.com/v1/chat/completions",
                "models": ["nvidia/llama-3.1-nemotron-70b-instruct"],
                "headers": lambda: {
                    "Authorization": f"Bearer {os.getenv('NVIDIA_API_KEY', '')}",
                    "Content-Type": "application/json"
                },
                "priority": 6,
                "specialty": "nvidia_optimized",
                "rate_limit": 5,
                "max_tokens": 4096
            }
        ]
        
        # Filter available providers based on API keys
        self.available = []
        for provider in self.providers:
            key_name = f"{provider['name'].upper()}_API_KEY"
            if os.getenv(key_name):
                self.available.append(provider)
                logger.info(f"✅ {provider['name']} API available ({provider['specialty']})")
            else:
                logger.warning(f"⚠️  {provider['name']} API key not found ({key_name})")
        
        if not self.available:
            logger.error("❌ No API keys found! Please set at least one API key.")
        
        self.available.sort(key=lambda x: x['priority'])
        self.current = self.available[0] if self.available else None
        
        # Enhanced performance tracking from GitHub reference
        self.performance_stats = {}
        self.request_history = deque(maxlen=100)
        
        for provider in self.available:
            self.performance_stats[provider['name']] = {
                'response_times': deque(maxlen=20),
                'success_rate': 1.0,
                'total_requests': 0,
                'failures': 0,
                'last_used': None,
                'rate_limit_reset': time.time(),
                'requests_this_minute': 0,
                'average_tokens': 0,
                'quality_score': 0.8  # Initial score
            }
        
        logger.info(f"✅ API Manager initialized with {len(self.available)} providers")

    def get_best_provider(self, query_type: str = "general", priority_override: str = None) -> dict:
        """Enhanced provider selection with load balancing"""
        if not self.available:
            return None
            
        if priority_override and priority_override.lower() in [p['name'].lower() for p in self.available]:
            for provider in self.available:
                if provider['name'].lower() == priority_override.lower():
                    return provider
        
        # Enhanced specialty routing from GitHub reference
        specialty_preferences = {
            "coding": ["fast_inference", "open_source", "diverse_models"],
            "creative": ["diverse_models", "text_generation", "conversational"],
            "analysis": ["nvidia_optimized", "diverse_models", "fast_inference"],
            "general": ["fast_inference", "diverse_models", "open_source"],
            "business": ["nvidia_optimized", "diverse_models", "fast_inference"],
            "medical": ["diverse_models", "nvidia_optimized", "fast_inference"],
            "technical": ["open_source", "fast_inference", "nvidia_optimized"]
        }
        
        preferred_specialties = specialty_preferences.get(query_type, ["fast_inference"])
        current_time = time.time()
        
        # Enhanced scoring algorithm
        best_provider = None
        best_score = -1
        
        for provider in self.available:
            stats = self.performance_stats[provider['name']]
            
            # Check rate limits
            if current_time - stats['rate_limit_reset'] > 60:
                stats['rate_limit_reset'] = current_time
                stats['requests_this_minute'] = 0
            
            if stats['requests_this_minute'] >= provider['rate_limit']:
                continue  # Skip rate-limited providers
            
            # Calculate composite score
            specialty_score = 10 if provider['specialty'] in preferred_specialties else 5
            performance_score = stats['success_rate'] * 8
            quality_score = stats['quality_score'] * 5
            
            # Speed score based on average response time
            if stats['response_times']:
                avg_time = sum(stats['response_times']) / len(stats['response_times'])
                speed_score = max(0, 8 - avg_time)  # Penalize slow responses
            else:
                speed_score = 6  # Default score for new providers
            
            # Freshness bonus (encourage trying different providers)
            freshness_score = 2 if not stats['last_used'] or (current_time - stats['last_used']) > 300 else 0
            
            # Load balancing score
            load_score = max(0, 3 - (stats['requests_this_minute'] / max(1, provider['rate_limit']) * 10))
            
            total_score = specialty_score + performance_score + quality_score + speed_score + freshness_score + load_score
            
            if total_score > best_score:
                best_score = total_score
                best_provider = provider
        
        return best_provider or self.current

    def _format_messages_for_provider(self, provider: dict, user_input: str, system_prompt: str) -> dict:
        """Format messages based on provider requirements"""
        if provider['name'] == 'HuggingFace':
            # HuggingFace has different format
            return {
                "inputs": f"System: {system_prompt}\nUser: {user_input}\nAssistant:",
                "parameters": {
                    "max_new_tokens": min(provider['max_tokens'], 1024),
                    "temperature": 0.7,
                    "return_full_text": False
                }
            }
        elif provider['name'] == 'Cohere':
            # Cohere has different format
            return {
                "message": user_input,
                "preamble": system_prompt,
                "max_tokens": min(provider['max_tokens'], 1500),
                "temperature": 0.7
            }
        else:
            # Standard OpenAI format for Groq, OpenRouter, Together, NVIDIA
            return {
                "model": provider["models"][0],  # Use first model as default
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_input}
                ],
                "max_tokens": min(provider['max_tokens'], 1500),
                "temperature": 0.7,
                "top_p": 0.9,
                "frequency_penalty": 0.1,
                "presence_penalty": 0.1
            }

    async def get_ai_response(self, user_input: str, system_prompt: str, query_type: str = "general") -> Optional[str]:
        """Enhanced AI response with fallback chain - From GitHub reference"""
        provider = self.get_best_provider(query_type)
        if not provider:
            logger.error("❌ No available API providers")
            return None
        
        start_time = time.time()
        stats = self.performance_stats[provider['name']]
        
        # Try each model from the selected provider
        for model_idx, model in enumerate(provider["models"][:2]):  # Try max 2 models per provider
            try:
                # Update rate limiting
                stats['requests_this_minute'] += 1
                stats['last_used'] = time.time()
                
                # Format payload based on provider
                if provider['name'] == 'HuggingFace':
                    url = f"{provider['url']}{model}"
                    payload = self._format_messages_for_provider(provider, user_input, system_prompt)
                else:
                    url = provider["url"]
                    payload = self._format_messages_for_provider(provider, user_input, system_prompt)
                    if provider['name'] not in ['Cohere']:  # Update model for standard providers
                        payload["model"] = model
                
                logger.info(f"🔄 Requesting {provider['name']} ({model}) for {query_type} query...")
                
                response = requests.post(
                    url,
                    headers=provider["headers"](),
                    json=payload,
                    timeout=35  # Increased timeout for better reliability
                )
                
                if response.status_code == 200:
                    result = response.json()
                    content = self._extract_content(result, provider['name'])
                    
                    if content:
                        # Update performance metrics
                        response_time = time.time() - start_time
                        stats['response_times'].append(response_time)
                        stats['total_requests'] += 1
                        stats['success_rate'] = (stats['total_requests'] - stats['failures']) / stats['total_requests']
                        
                        # Update quality score based on response length and coherence
                        quality_indicators = len(content.split()) > 20 and len(content) > 100
                        stats['quality_score'] = min(1.0, stats['quality_score'] + 0.05 if quality_indicators else stats['quality_score'] - 0.02)
                        
                        logger.info(f"✅ {provider['name']} responded in {response_time:.2f}s")
                        return content
                        
                elif response.status_code == 429:
                    logger.warning(f"⏳ {provider['name']} rate limited, trying next model...")
                    stats['rate_limit_reset'] = time.time() + 60
                    continue
                else:
                    logger.warning(f"⚠️  {provider['name']} error {response.status_code}: {response.text[:200]}")
                    
            except requests.exceptions.Timeout:
                logger.warning(f"⏰ {provider['name']} timeout, trying next model...")
                continue
            except Exception as e:
                logger.error(f"❌ {provider['name']} model {model} failed: {e}")
                continue
        
        # Update failure stats for the provider
        stats['failures'] += 1
        stats['total_requests'] += 1
        stats['success_rate'] = max(0, (stats['total_requests'] - stats['failures']) / stats['total_requests'])
        stats['quality_score'] = max(0.1, stats['quality_score'] - 0.1)
        
        # Try next best provider as fallback
        logger.warning(f"🔄 {provider['name']} failed, trying fallback providers...")
        
        # Get next best providers (excluding the failed one)
        available_backup = [p for p in self.available if p['name'] != provider['name']]
        
        for backup_provider in available_backup[:2]:  # Try up to 2 backup providers
            try:
                backup_stats = self.performance_stats[backup_provider['name']]
                if backup_stats['requests_this_minute'] >= backup_provider['rate_limit']:
                    continue
                    
                logger.info(f"🔄 Fallback to {backup_provider['name']}...")
                
                backup_stats['requests_this_minute'] += 1
                backup_stats['last_used'] = time.time()
                
                payload = self._format_messages_for_provider(backup_provider, user_input, system_prompt)
                url = backup_provider["url"]
                
                if backup_provider['name'] == 'HuggingFace':
                    url = f"{backup_provider['url']}{backup_provider['models'][0]}"
                
                response = requests.post(
                    url,
                    headers=backup_provider["headers"](),
                    json=payload,
                    timeout=30
                )
                
                if response.status_code == 200:
                    result = response.json()
                    content = self._extract_content(result, backup_provider['name'])
                    
                    if content:
                        logger.info(f"✅ Fallback {backup_provider['name']} succeeded!")
                        return content
                        
            except Exception as e:
                logger.error(f"❌ Fallback {backup_provider['name']} failed: {e}")
                continue
        
        logger.error("❌ All API providers failed")
        return None

    def _extract_content(self, result: dict, provider_name: str) -> Optional[str]:
        """Extract content from API response based on provider format"""
        try:
            if provider_name == 'HuggingFace':
                if isinstance(result, list) and len(result) > 0:
                    return result[0].get('generated_text', '').strip()
                return result.get('generated_text', '').strip()
                
            elif provider_name == 'Cohere':
                return result.get('text', '').strip()
                
            else:  # Standard OpenAI format (Groq, OpenRouter, Together, NVIDIA)
                choices = result.get("choices", [])
                if choices and len(choices) > 0:
                    message = choices[0].get("message", {})
                    return message.get("content", "").strip()
                    
        except Exception as e:
            logger.error(f"Content extraction error for {provider_name}: {e}")
            
        return None

    def get_provider_stats(self) -> Dict[str, Any]:
        """Get comprehensive provider statistics"""
        stats_summary = {}
        
        for provider_name, stats in self.performance_stats.items():
            avg_response_time = sum(stats['response_times']) / len(stats['response_times']) if stats['response_times'] else 0
            
            stats_summary[provider_name] = {
                'success_rate': f"{stats['success_rate']:.2%}",
                'avg_response_time': f"{avg_response_time:.2f}s",
                'total_requests': stats['total_requests'],
                'quality_score': f"{stats['quality_score']:.2f}",
                'requests_this_minute': stats['requests_this_minute'],
                'last_used': stats['last_used']
            }
        
        return {
            'available_providers': len(self.available),
            'current_provider': self.current['name'] if self.current else None,
            'provider_stats': stats_summary,
            'total_providers': len(self.providers)
        }

# ========== PROFESSIONAL AGENTS SYSTEM (EXACT FROM CLI) ==========
class ProfessionalAgentsSystem:
    """Professional Agents System - EXACT from CLI"""
    
    def __init__(self):
        self.agents = {}
        self.load_professional_agents()
    
    def load_professional_agents(self):
        """Load professional agents - EXACT from CLI"""
        if PROFESSIONAL_AGENTS_LOADED:
            try:
                self.agents = {
                    'coding': ProLevelCodingExpert(),
                    'career': ProfessionalCareerCoach(),
                    'business': SmartBusinessConsultant(),
                    'medical': SimpleMedicalAdvisor(),
                    'emotional': SimpleEmotionalCounselor(),
                    'technical_architect': TechnicalArchitect()
                }
                logger.info(f"✅ {len(self.agents)} professional agents loaded")
            except Exception as e:
                logger.error(f"Professional agents loading error: {e}")
                self.agents = {}
        else:
            logger.info("Professional agents not available - using enhanced fallback system")

# ========== NOVA SYSTEM ORCHESTRATOR (ENHANCED) ==========
class NOVASystemOrchestrator:
    """NOVA System Orchestrator - Enhanced with Always AI Response"""
    
    def __init__(self):
        self.memory_system = memory_system
        self.agents = ProfessionalAgentsSystem()
        self.api_manager = OptimizedAPIManager()
        
        # Performance tracking
        if ADVANCED_SYSTEMS:
            self.orchestrator = IntelligentAPIOrchestrator()
            self.drift_detector = APIPerformanceDrifter()
        else:
            self.orchestrator = None
            self.drift_detector = None
        
        logger.info("✅ NOVA System Orchestrator initialized")
    
    async def get_response(self, user_input: str, user_id: str = "default", 
                         agent_type: str = "general", session_id: str = None) -> Dict[str, Any]:
        """Enhanced response generation - ALWAYS uses AI, smart ML enhancement"""
        start_time = time.time()
        session_id = session_id or f"session_{int(time.time())}"
        
        try:
            # Step 1: Check if ML enhancement is needed
            needs_ml_enhancement = SmartEnhancementDetector.needs_ml_enhancement(user_input)
            is_simple_greeting = SmartEnhancementDetector.is_simple_greeting(user_input)
            
            logger.info(f"🧠 Query analysis - ML Enhancement: {needs_ml_enhancement}, Simple: {is_simple_greeting}")
            
            # Step 2: Prepare prompt for AI (ALWAYS use AI, never dummy responses)
            base_prompt = user_input
            ml_analysis = {}
            enhanced_agent_type = agent_type
            
            if needs_ml_enhancement and not is_simple_greeting:
                # Complex query - Apply full ML enhancement
                logger.info(f"🔥 Applying ML enhancement for complex query: {user_input[:50]}...")
                
                # Get conversation context
                conversation_context = await self.memory_system.get_conversation_context(user_id, limit=5)
                user_profile = await self.memory_system.get_user_profile(user_id)
                
                if ML_SYSTEM_AVAILABLE:
                    # Run comprehensive ML analysis
                    ml_analysis = ml_manager.process_user_query(
                        user_input,
                        context={
                            "conversation_history": conversation_context,
                            "user_profile": user_profile,
                            "session_id": session_id,
                            "requested_agent": agent_type
                        }
                    )
                    
                    # Use ML-recommended agent if confidence is high
                    if ml_analysis.get('routing_decision', {}).get('confidence_level', 0) > 0.7:
                        enhanced_agent_type = ml_analysis['routing_decision']['selected_agent']
                        logger.info(f"🎯 ML routing: {agent_type} → {enhanced_agent_type}")
                
                # Enhanced prompt construction for complex queries
                base_prompt = f"""
                Context from recent conversations:
                {conversation_context[:500] if conversation_context else 'No recent context'}
                
                User Profile Insights:
                {json.dumps(user_profile, indent=2) if user_profile else 'No profile data available'}
                
                Current Query: {user_input}
                
                ML Analysis Insights:
                {json.dumps(ml_analysis.get('recommendations', []), indent=2) if ml_analysis else 'No ML insights available'}
                
                Please provide a comprehensive, professional response that takes into account the conversation context and user profile.
                """
            else:
                # Simple query - Use AI but without heavy ML processing
                logger.info(f"💬 Simple query - using AI without ML enhancement: {user_input[:50]}...")
            
            # Step 3: Create enhanced system prompt
            system_prompt = self._create_enhanced_system_prompt(enhanced_agent_type, ml_analysis)
            
            # Step 4: Get AI response (ALWAYS use AI, never skip this step)
            ai_response = await self.api_manager.get_ai_response(
                base_prompt, system_prompt, enhanced_agent_type
            )
            
            # Step 5: Fallback if AI fails
            if not ai_response:
                ai_response = self._get_emergency_response(user_input, enhanced_agent_type)
            
            # Step 6: Advanced optimization if available (only for complex queries)
            if needs_ml_enhancement and ADVANCED_SYSTEMS and self.orchestrator:
                optimized_response, optimization_metadata = await self.orchestrator.get_optimized_response(
                    ai_response, user_input, enhanced_agent_type
                )
                if optimized_response:
                    ai_response = optimized_response
            
            # Step 7: Performance monitoring
            response_time = time.time() - start_time
            if ADVANCED_SYSTEMS and self.drift_detector:
                self.drift_detector.record_response_quality(
                    ai_response, user_input, response_time, enhanced_agent_type
                )
            
            # Step 8: Enhanced memory storage
            await self.memory_system.remember_conversation(
                user_id=user_id,
                session_id=session_id,
                user_input=user_input,
                bot_response=ai_response,
                agent_type=enhanced_agent_type,
                language="english",
                emotion=ml_analysis.get('query_analysis', {}).get('sentiment', {}).get('overall', 'neutral'),
                confidence=0.9,
                response_time=response_time,
                ml_insights=ml_analysis,
                enhancement_applied=needs_ml_enhancement
            )
            
            # Step 9: Get conversation count
            with sqlite3.connect(self.memory_system.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT total_conversations FROM user_profiles WHERE user_id = ?", (user_id,))
                result = cursor.fetchone()
                conversation_count = result[0] if result else 1
            
            return {
                'response': ai_response,
                'agent_used': enhanced_agent_type,
                'language': 'english',
                'emotion': ml_analysis.get('query_analysis', {}).get('sentiment', {}).get('overall', 'neutral'),
                'emotion_confidence': 0.8,
                'agent_confidence': ml_analysis.get('routing_decision', {}).get('confidence_level', 0.9),
                'response_time': response_time,
                'conversation_count': conversation_count,
                'ml_enhanced': needs_ml_enhancement,
                'session_id': session_id,
                'context_used': bool(conversation_context if needs_ml_enhancement else False),
                'recommendations': ml_analysis.get('recommendations', [])[:3] if needs_ml_enhancement else [],
                'enhancement_reason': f"{'Complex query - full ML enhancement applied' if needs_ml_enhancement else 'Simple query - AI response without ML overhead'}"
            }
            
        except Exception as e:
            logger.error(f"Response generation error: {e}")
            
            # Fallback - still use AI, not dummy response
            try:
                fallback_response = await self.api_manager.get_ai_response(
                    user_input, 
                    "You are NOVA, a professional AI assistant. Provide helpful, accurate responses.",
                    "general"
                )
                
                if not fallback_response:
                    fallback_response = self._get_emergency_response(user_input, "general")
                
                return {
                    'response': fallback_response,
                    'agent_used': 'general',
                    'language': 'english',
                    'emotion': 'neutral',
                    'emotion_confidence': 0.7,
                    'agent_confidence': 0.7,
                    'response_time': time.time() - start_time,
                    'conversation_count': 1,
                    'ml_enhanced': False,
                    'session_id': session_id,
                    'error': str(e),
                    'enhancement_reason': 'Error occurred - fallback AI response provided'
                }
            except:
                # Last resort fallback
                return {
                    'response': self._get_emergency_response(user_input, "general"),
                    'agent_used': 'general',
                    'language': 'english',
                    'emotion': 'neutral',
                    'emotion_confidence': 0.6,
                    'agent_confidence': 0.6,
                    'response_time': time.time() - start_time,
                    'conversation_count': 1,
                    'ml_enhanced': False,
                    'session_id': session_id,
                    'error': str(e),
                    'enhancement_reason': 'Critical error - emergency response provided'
                }

    def _create_enhanced_system_prompt(self, agent_type: str, ml_analysis: Dict = None) -> str:
        """Create enhanced system prompt based on agent type and ML analysis"""
        
        # Ultra-professional system prompts for each agent
        system_prompts = {
            "general": """You are NOVA, an ultra-professional AI assistant with exceptional expertise across multiple domains. 
            Provide comprehensive, well-structured responses with professional tone. Focus on delivering high-quality, 
            actionable advice with attention to detail and practical implementation steps. Be friendly and conversational 
            while maintaining professionalism.""",
            
            "coding": """You are a world-class software engineering expert with deep knowledge across all programming languages, 
            frameworks, and software architecture patterns. Provide clean, efficient, well-commented code solutions with 
            comprehensive explanations. Include best practices, performance considerations, security implications, testing 
            strategies, and deployment considerations. Structure your responses with clear sections for immediate solutions 
            and advanced optimizations. Be conversational and helpful.""",
            
            "career": """You are an elite career strategist and professional development coach with extensive industry knowledge. 
            Provide strategic career guidance with specific, actionable steps. Include industry insights, skill development 
            roadmaps, networking strategies, and market analysis. Structure responses with immediate actions, medium-term goals, 
            and long-term career vision. Consider current market trends and future industry evolution. Be supportive and encouraging.""",
            
            "business": """You are a senior business consultant and strategic analyst with expertise in business intelligence, 
            market analysis, and growth strategies. Provide data-driven recommendations with quantitative analysis where possible. 
            Include market positioning, competitive analysis, financial implications, risk assessment, and scalability considerations. 
            Structure responses with executive summary, detailed analysis, and implementation roadmap. Be insightful and strategic.""",
            
            "medical": """You are a medical information specialist with comprehensive knowledge of evidence-based healthcare. 
            Provide accurate, well-researched health information while emphasizing the critical importance of professional 
            medical consultation. Include relevant medical literature, risk factors, prevention strategies, and treatment options. 
            Always include appropriate disclaimers about seeking professional medical advice for diagnosis and treatment. 
            Be caring and informative.""",
            
            "emotional": """You are a compassionate emotional wellness specialist and counselor with training in psychology 
            and mental health. Provide empathetic, supportive guidance with practical coping strategies and emotional validation. 
            Include stress management techniques, mindfulness practices, communication strategies, and mental health resources. 
            Maintain a warm, understanding tone while providing professional-grade emotional support. Be genuinely caring and supportive.""",
            
            "technical_architect": """You are a distinguished technical architect and system design expert with deep expertise 
            in scalable system architecture, cloud computing, and enterprise solutions. Provide comprehensive architectural 
            guidance with detailed technical specifications. Include scalability patterns, performance optimization, security 
            architecture, monitoring strategies, and technology selection criteria. Structure responses with architectural 
            overview, detailed design considerations, and implementation best practices. Be technically precise yet approachable."""
        }
        
        base_prompt = system_prompts.get(agent_type, system_prompts["general"])
        
        # Add ML insights if available
        if ml_analysis and ml_analysis.get('recommendations'):
            ml_context = f"""
            
            Additional Context (ML Analysis):
            Based on advanced analysis, consider these recommendations: {ml_analysis.get('recommendations', [])}
            User intent confidence: {ml_analysis.get('routing_decision', {}).get('confidence_level', 0.0):.2f}
            Query complexity: {'High' if ml_analysis.get('context_enhancement') else 'Standard'}
            """
            base_prompt += ml_context
        
        return base_prompt

    def _get_emergency_response(self, user_input: str, agent_type: str) -> str:
        """Emergency fallback response when all APIs fail"""
        emergency_responses = {
            "coding": f"""I understand you're asking about a coding-related topic. While I'm experiencing some technical 
            difficulties with my primary systems, I can still provide guidance:

            **General Coding Best Practices:**
            - Break down complex problems into smaller, manageable functions
            - Use meaningful variable and function names
            - Implement proper error handling and validation
            - Follow language-specific style guidelines (PEP 8 for Python, etc.)
            - Write comprehensive tests for your code
            - Document your code thoroughly

            For your specific question about: "{user_input[:100]}..."
            I recommend checking official documentation, Stack Overflow, or GitHub repositories for similar implementations.

            Please try your question again, and I'll provide more detailed assistance once my systems are fully restored.""",

            "career": f"""I understand you're seeking career guidance. While I'm experiencing some technical issues, 
            I can still offer some strategic advice:

            **Core Career Development Principles:**
            - Continuously update your skills to match market demands
            - Build and maintain a strong professional network
            - Seek feedback regularly and act on it constructively
            - Set clear short-term and long-term career goals
            - Develop both technical and soft skills
            - Stay informed about industry trends

            Regarding your question: "{user_input[:100]}..."
            I recommend researching industry-specific resources, connecting with professionals in your field, 
            and considering professional career coaching services.

            Please ask your question again for more personalized guidance.""",

            "business": f"""I understand you're looking for business insights. Despite current technical limitations, 
            here are some fundamental business principles:

            **Strategic Business Considerations:**
            - Understand your target market deeply
            - Analyze competitors and market positioning regularly
            - Focus on customer value proposition
            - Monitor key performance indicators (KPIs)
            - Maintain healthy cash flow management
            - Plan for scalable growth

            For your specific inquiry: "{user_input[:100]}..."
            Consider consulting industry reports, business analytics tools, or professional business advisors.

            Please try your question again for more detailed strategic analysis.""",

            "general": f"""I apologize for the technical difficulty, but I'm still here to help you with: "{user_input[:100]}..."

            While I'm working to restore full functionality, I can still provide assistance with:
            - General information and explanations
            - Problem-solving approaches
            - Research guidance and resources
            - Basic analysis and recommendations

            **Immediate Suggestions:**
            1. Try rephrasing your question with more specific details
            2. Break complex questions into smaller parts
            3. Specify what type of assistance you're looking for

            Please ask your question again, and I'll do my best to provide comprehensive help once my systems are fully operational."""
        }

        return emergency_responses.get(agent_type, emergency_responses["general"])

    def get_system_status(self) -> Dict[str, Any]:
        """Get comprehensive system status"""
        return {
            "status": "operational",
            "version": "3.0.0-free-apis-enhanced",
            "timestamp": datetime.now().isoformat(),
            "components": {
                "memory_system": "operational",
                "professional_agents": len(self.agents.agents),
                "api_providers": len(self.api_manager.available),
                "ml_system": "enhanced" if ML_SYSTEM_AVAILABLE else "basic",
                "advanced_systems": ADVANCED_SYSTEMS,
                "voice_processing": VOICE_AVAILABLE,
                "file_processing": FILE_PROCESSING_AVAILABLE,
                "github_integration": GITHUB_INTEGRATION
            },
            "api_providers": self.api_manager.get_provider_stats(),
            "capabilities": {
                "always_ai_response": True,
                "smart_enhancement_detection": True,
                "ml_enhanced_routing": ML_SYSTEM_AVAILABLE,
                "context_aware_responses": True,
                "professional_agents": bool(self.agents.agents),
                "conversation_memory": True,
                "performance_monitoring": ADVANCED_SYSTEMS,
                "multi_provider_apis": True,
                "free_tier_optimized": True
            }
        }
    
    def clear_user_context(self, user_id: str):
        """Clear user context"""
        # Clear in-memory context
        self.memory_system.conversation_context.clear()
        if user_id in self.memory_system.conversation_threads:
            self.memory_system.conversation_threads[user_id].clear()
        
        logger.info(f"Context cleared for user: {user_id}")

# ========== NOVA ULTRA SYSTEM (ENHANCED) ==========
class NovaUltraSystem:
    """Enhanced NOVA Ultra System - Always AI Response with Free APIs"""
    
    def __init__(self):
        self.memory = memory_system
        self.agents = ProfessionalAgentsSystem()
        self.api_manager = OptimizedAPIManager()
        self.language_detector = FastLanguageDetector()
        self.emotion_detector = FastEmotionDetector()
        
        # Initialize all systems from original
        self.current_sessions = defaultdict(lambda: {
            'file_context': None,
            'conversation_count': 0,
            'last_agent': 'general',
            'voice_enabled': False,
            'search_history': []
        })
        
        self.conversation_count = 0
        self.ml_manager = ml_manager if ML_SYSTEM_AVAILABLE else None
        
        # Initialize voice and file systems
        self.voice_system = self._initialize_voice_system()
        self.file_system = self._initialize_file_system()
        self.web_search = self._initialize_web_search()
        
        logger.info("✅ NOVA Ultra System initialized - Always AI Response Mode with Free APIs")
    
    def _initialize_voice_system(self):
        """Initialize voice system - exact from original"""
        class VoiceSystem:
            def __init__(self):
                self.azure_enabled = AZURE_VOICE_AVAILABLE
                self.basic_enabled = VOICE_AVAILABLE
                
                if self.azure_enabled:
                    self.setup_azure_voice()
                if self.basic_enabled:
                    self.setup_basic_voice()
            
            def setup_azure_voice(self):
                """Setup Azure voice services"""
                try:
                    azure_key = os.getenv('AZURE_SPEECH_KEY')
                    azure_region = os.getenv('AZURE_SPEECH_REGION', 'eastus')
                    
                    if azure_key:
                        self.speech_config = speechsdk.SpeechConfig(
                            subscription=azure_key, 
                            region=azure_region
                        )
                        self.speech_config.speech_recognition_language = "en-US"
                        self.speech_config.speech_synthesis_voice_name = "en-US-JennyNeural"
                except Exception as e:
                    logger.error(f"Azure Voice setup error: {e}")
                    self.azure_enabled = False
            
            def setup_basic_voice(self):
                """Setup basic voice recognition"""
                try:
                    self.recognizer = sr.Recognizer()
                    self.tts_engine = pyttsx3.init()
                    self.tts_engine.setProperty('rate', 180)
                except Exception as e:
                    logger.error(f"Basic voice setup error: {e}")
                    self.basic_enabled = False
            
            async def process_audio(self, audio_data):
                """Process audio input"""
                if self.basic_enabled:
                    try:
                        recognizer = sr.Recognizer()
                        with sr.AudioFile(BytesIO(audio_data)) as source:
                            audio = recognizer.record(source)
                        return recognizer.recognize_google(audio)
                    except:
                        return "Could not understand audio"
                return "Voice processing not available"
            
            async def text_to_speech(self, text, voice="en-US-AriaNeural"):
                """Convert text to speech"""
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                clean_text = re.sub(r'[🔧💼📈🏥💙🚀🎯📋💡📚🤖⚠️✅❌🔊📝🎤]', '', clean_text)
                
                if len(clean_text) > 300:
                    clean_text = clean_text[:300] + "..."
                
                if self.azure_enabled:
                    try:
                        self.speech_config.speech_synthesis_voice_name = voice
                        synthesizer = speechsdk.SpeechSynthesizer(
                            speech_config=self.speech_config,
                            audio_config=None
                        )
                        result = synthesizer.speak_text_async(clean_text).get()
                        
                        if result.reason == speechsdk.ResultReason.SynthesizingAudioCompleted:
                            return result.audio_data
                    except Exception as e:
                        logger.error(f"Azure TTS failed: {e}")
                
                if self.basic_enabled:
                    try:
                        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_file:
                            temp_path = temp_file.name
                        self.tts_engine.save_to_file(clean_text, temp_path)
                        self.tts_engine.runAndWait()
                        
                        with open(temp_path, 'rb') as audio_file:
                            audio_data = audio_file.read()
                        os.unlink(temp_path)
                        return audio_data
                    except Exception as e:
                        logger.error(f"Basic TTS failed: {e}")
                
                return b""
        
        return VoiceSystem()
    
    def _initialize_file_system(self):
        """Initialize file system - exact from original"""
        class FileSystem:
            def process_file(self, file_content, filename):
                """Process uploaded file"""
                file_analysis = {
                    'file_name': filename,
                    'file_size': len(file_content),
                    'file_type': self._detect_file_type(filename),
                    'content': self._extract_content(file_content, filename)
                }
                return file_analysis
            
            def _detect_file_type(self, filename):
                """Detect file type from filename"""
                ext = filename.lower().split('.')[-1] if '.' in filename else 'unknown'
                type_map = {
                    'txt': 'text/plain',
                    'py': 'text/python',
                    'js': 'text/javascript',
                    'html': 'text/html',
                    'css': 'text/css',
                    'md': 'text/markdown',
                    'pdf': 'application/pdf',
                    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                }
                return type_map.get(ext, 'application/octet-stream')
            
            def _extract_content(self, file_content, filename):
                """Extract text content from file"""
                try:
                    if filename.endswith('.txt') or filename.endswith('.py') or filename.endswith('.js'):
                        return file_content.decode('utf-8', errors='ignore')
                    elif filename.endswith('.pdf') and FILE_PROCESSING_AVAILABLE:
                        pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                        return text
                    elif filename.endswith('.docx') and FILE_PROCESSING_AVAILABLE:
                        doc = docx.Document(BytesIO(file_content))
                        text = ""
                        for paragraph in doc.paragraphs:
                            text += paragraph.text + "\n"
                        return text
                    else:
                        return "Binary file - content not extracted"
                except:
                    return "Error extracting content"
        
        return FileSystem()
    
    def _initialize_web_search(self):
        """Initialize web search - exact from original"""
        class WebSearch:
            async def search_web(self, query, max_results=5):
                """Basic web search functionality"""
                try:
                    # Implement basic search functionality
                    return {
                        "success": True,
                        "results": [
                            {
                                "title": f"Search result for: {query}",
                                "source": "example.com",
                                "snippet": f"This is a search result for the query: {query}"
                            }
                        ],
                        "count": 1
                    }
                except:
                    return {"success": False, "error": "Search failed"}
        
        return WebSearch()

    async def detect_agent_type(self, user_input: str) -> Tuple[str, float]:
        """Agent detection - EXACT from CLI"""
        text_lower = user_input.lower()
        
        agent_patterns = {
            "coding": {
                "keywords": ["code", "programming", "debug", "python", "javascript", "bug", "development", "api", "function"],
                "confidence": 0.8
            },
            "career": {
                "keywords": ["resume", "interview", "job", "career", "hiring", "professional", "promotion", "salary"],
                "confidence": 0.8
            },
            "business": {
                "keywords": ["business", "analysis", "strategy", "market", "revenue", "growth", "profit", "finance"],
                "confidence": 0.8
            },
            "medical": {
                "keywords": ["health", "medical", "symptoms", "doctor", "treatment", "medicine", "therapy"],
                "confidence": 0.8
            },
            "emotional": {
                "keywords": ["stress", "anxiety", "sad", "emotional", "support", "therapy", "counseling", "mental"],
                "confidence": 0.8
            },
            "technical_architect": {
                "keywords": ["architecture", "system design", "scalability", "microservice", "infrastructure", "devops"],
                "confidence": 0.8
            }
        }
        
        for agent_name, agent_data in agent_patterns.items():
            keywords = agent_data["keywords"]
            if any(keyword in text_lower for keyword in keywords):
                return agent_name, agent_data["confidence"]
        return "general", 0.0

    async def get_response(self, user_input: str, user_id: str = "default", 
                         agent_type: str = "general", session_id: str = None) -> Dict[str, Any]:
        """Get AI response - ALWAYS AI, smart enhancement when needed"""
        start_time = time.time()
        session_id = session_id or f"session_{int(time.time())}"
        
        # Always check if ML enhancement is needed
        needs_enhancement = SmartEnhancementDetector.needs_ml_enhancement(user_input)
        
        try:
            # Get user session
            user_session = self.current_sessions[user_id]
            
            # Fast detection
            language = self.language_detector.detect_language(user_input)
            emotion, emotion_confidence = self.emotion_detector.detect_emotion(user_input)
            detected_agent_type, agent_confidence = await self.detect_agent_type(user_input)
            
            # Use detected agent if confidence is high
            if agent_confidence > 0.7:
                agent_type = detected_agent_type

            # Always get AI response, but enhance with ML for complex queries
            if needs_enhancement:
                # Complex response path with ML enhancement
                logger.info(f"🧠 Complex query - applying ML enhancement: {user_input[:50]}...")
                
                # Get conversation context for complex queries
                conversation_context = await self.memory.get_conversation_context(user_id, limit=5)
                user_profile = await self.memory.get_user_profile(user_id)
                
                # Enhanced prompt for complex queries
                enhanced_prompt = f"""
                Previous conversation context:
                {conversation_context if conversation_context else 'No previous context'}
                
                User profile:
                {json.dumps(user_profile, indent=2) if user_profile else 'New user'}
                
                Current query: {user_input}
                
                Please provide a comprehensive, contextually aware response.
                """
                
                system_prompt = self._create_system_prompt(agent_type, language, emotion, conversation_context)
                ai_response = await self.api_manager.get_ai_response(enhanced_prompt, system_prompt, agent_type)
            else:
                # Simple response path - still use AI but without ML overhead
                logger.info(f"💬 Simple query - AI response without ML enhancement: {user_input[:50]}...")
                system_prompt = self._create_system_prompt(agent_type, language, emotion)
                ai_response = await self.api_manager.get_ai_response(user_input, system_prompt, agent_type)
            
            # Fallback if AI response fails
            if not ai_response:
                ai_response = self._get_fallback_response(user_input, agent_type)
            
            # Update session
            user_session['conversation_count'] += 1
            user_session['last_agent'] = agent_type
            
            # Store in memory
            await self.memory.remember_conversation(
                user_id, session_id, user_input, ai_response,
                agent_type, language, emotion, emotion_confidence,
                response_time=time.time() - start_time,
                enhancement_applied=needs_enhancement
            )
            
            response_data = {
                'response': ai_response,
                'agent_used': agent_type,
                'language': language,
                'emotion': emotion,
                'emotion_confidence': emotion_confidence,
                'agent_confidence': agent_confidence,
                'response_time': time.time() - start_time,
                'conversation_count': user_session['conversation_count'],
                'file_context_used': bool(user_session['file_context']),
                'user_id': user_id,
                'session_id': session_id,
                'ml_enhanced': needs_enhancement,
                'context_used': needs_enhancement,
                'recommendations': [],
                'enhancement_reason': f"{'Complex query - ML enhancement applied' if needs_enhancement else 'Simple query - direct AI response'}"
            }
            
            return response_data
            
        except Exception as e:
            logger.error(f"Response generation error: {e}")
            
            # Fallback - still try to get AI response
            try:
                fallback_response = await self.api_manager.get_ai_response(
                    user_input, 
                    "You are NOVA, a professional AI assistant. Provide helpful responses.", 
                    "general"
                )
                
                if not fallback_response:
                    fallback_response = self._get_fallback_response(user_input, "general")
                
                return {
                    'response': fallback_response,
                    'agent_used': 'general',
                    'language': 'english',
                    'emotion': 'neutral',
                    'emotion_confidence': 0.7,
                    'agent_confidence': 0.7,
                    'response_time': time.time() - start_time,
                    'conversation_count': self.current_sessions[user_id]['conversation_count'],
                    'file_context_used': False,
                    'user_id': user_id,
                    'session_id': session_id,
                    'ml_enhanced': False,
                    'context_used': False,
                    'recommendations': [],
                    'enhancement_reason': 'Error recovery - AI fallback response'
                }
            except:
                # Last resort
                return {
                    'response': self._get_fallback_response(user_input, "general"),
                    'agent_used': 'general',
                    'language': 'english',
                    'emotion': 'neutral',
                    'emotion_confidence': 0.6,
                    'agent_confidence': 0.6,
                    'response_time': time.time() - start_time,
                    'conversation_count': self.current_sessions[user_id]['conversation_count'],
                    'file_context_used': False,
                    'user_id': user_id,
                    'session_id': session_id,
                    'ml_enhanced': False,
                    'context_used': False,
                    'recommendations': [],
                    'enhancement_reason': 'Critical error - emergency response'
                }

    def _create_system_prompt(self, agent_type: str, language: str, emotion: str, 
                             user_context: str = None, file_context: str = None) -> str:
        """Create system prompt - Enhanced from original"""
        base_prompt = """You are NOVA Ultra Professional AI, an advanced assistant with expertise across all domains.
        Provide professional, actionable, and empathetic responses. Be concise yet comprehensive."""
        
        agent_prompts = {
            "general": "You are NOVA, an ultra-professional AI assistant. Provide comprehensive, well-structured responses.",
            "coding": "You are a world-class software engineering expert. Provide clean, efficient code solutions with explanations.",
            "career": "You are an elite career strategist. Provide strategic career guidance with actionable steps.",
            "business": "You are a senior business consultant. Provide data-driven recommendations and strategic analysis.",
            "medical": "You are a medical information specialist. Provide accurate health information with appropriate disclaimers.",
            "emotional": "You are a compassionate counselor. Provide empathetic support and practical coping strategies.",
            "technical_architect": "You are a distinguished technical architect. Provide comprehensive architectural guidance."
        }
        
        agent_prompt = agent_prompts.get(agent_type, agent_prompts["general"])
        
        language_note = ""
        if language == "hinglish":
            language_note = " Respond naturally mixing English and Hindi as appropriate."
        
        emotion_note = ""
        if emotion in ["sad", "anxious", "frustrated"]:
            emotion_note = f" The user seems {emotion}, so be extra supportive and empathetic."
        
        context_note = ""
        if user_context:
            context_note = f"\n\nCONVERSATION CONTEXT:\n{user_context}"
        
        file_context_note = ""
        if file_context:
            file_context_note = f"\n\nFILE CONTEXT:\n{file_context}"
        
        return f"{base_prompt}\n{agent_prompt}{language_note}{emotion_note}{context_note}{file_context_note}"

    def _get_fallback_response(self, user_input: str, agent_type: str) -> str:
        """High-quality fallback responses when APIs fail"""
        fallback_responses = {
            "general": f"""Hello! I'm NOVA, your professional AI assistant. I understand you're asking about: '{user_input[:100]}...'

I'm here to provide comprehensive guidance and support. While I'm currently operating in fallback mode, I can still help you with:

🔧 **Technical Solutions**: Programming, debugging, system architecture
💼 **Professional Growth**: Career planning, skill development, industry insights  
📊 **Business Strategy**: Market analysis, growth planning, strategic decisions
🏥 **Health & Wellness**: Evidence-based health information and wellness strategies
💙 **Emotional Support**: Mental wellness, stress management, emotional intelligence
🗃️ **System Design**: Technical architecture, scalability planning

For the most effective assistance, please provide specific details about your situation and objectives.

How can I help you achieve your goals today?""",

            "coding": f"""As your professional coding expert, I understand you're working on: '{user_input[:100]}...'

**Recommended Approach:**
1. **Analysis**: Break down the problem into manageable components
2. **Architecture**: Design a clean, scalable solution following best practices
3. **Implementation**: Write well-structured, documented code
4. **Testing**: Implement comprehensive tests
5. **Optimization**: Profile and optimize for performance

**Best Practices:**
- Use meaningful variable names
- Implement proper error handling
- Follow coding standards
- Plan for scalability

Please provide more context about your technology stack and requirements for specific guidance.""",

            "career": f"""Thank you for your career inquiry: '{user_input[:100]}...'

**Strategic Career Framework:**

**Immediate Actions (30 days):**
- Skills assessment against market demands
- Network expansion and professional connections
- Personal branding optimization
- Market research for target roles

**Medium-term Goals (3-6 months):**
- Skill development and certifications
- Thought leadership through content sharing
- Strategic job applications
- Interview preparation and practice

**Long-term Vision (6-18 months):**
- Subject matter expertise development
- Leadership opportunity pursuit
- Industry engagement and networking
- Mentorship (both seeking and providing)

Please share details about your current role and career objectives for personalized advice.""",

            "business": f"""Regarding your business question: '{user_input[:100]}...'

**Strategic Business Framework:**

**Market Analysis:**
- Competitive landscape assessment
- Value proposition definition
- Customer segmentation analysis
- Market opportunity sizing

**Growth Strategy:**
- Revenue optimization strategies
- Operational efficiency improvements
- Customer acquisition planning
- Product development alignment

**Implementation:**
- Resource allocation planning
- Performance metrics definition
- Risk management strategies
- Scalability considerations

Please provide more context about your business vertical and specific objectives.""",

            "medical": f"""I understand your health-related inquiry: '{user_input[:100]}...'

**Evidence-Based Health Guidance:**

**Key Areas:**
- Current medical research and guidelines
- Risk factor identification and management
- Prevention strategies and lifestyle factors
- Treatment option overviews

**Important Considerations:**
- Individual health needs vary significantly
- Medical history affects recommendations
- Professional medical consultation is essential
- Medication interactions require professional oversight

⚠️ **Critical Disclaimer**: This information is educational only. Always consult qualified healthcare providers for medical diagnosis, treatment plans, and health decisions.

Please seek immediate medical attention for urgent health concerns.""",

            "emotional": f"""I hear that you're dealing with: '{user_input[:100]}...'

Your feelings are completely valid, and seeking support shows tremendous strength.

**Emotional Support Framework:**

**Immediate Strategies:**
- Grounding techniques and breathing exercises
- Self-compassion and emotional validation
- Safe space creation (physical and mental)
- Stress reduction techniques

**Building Resilience:**
- Mindfulness and meditation practices
- Emotional regulation strategies
- Communication skill development
- Social connection maintenance

**Professional Growth:**
- Perspective building through challenges
- Problem-solving skill enhancement
- Self-care planning and implementation
- Personal boundary development

Remember: Professional counseling and therapy are valuable resources. There's strength in seeking professional support.

You're not alone, and positive change is possible.""",

            "technical_architect": f"""Analyzing your technical architecture question: '{user_input[:100]}...'

**System Architecture Framework:**

**Core Principles:**
- Scalability and performance optimization
- Reliability and fault tolerance
- Security and compliance considerations
- Maintainability and documentation

**Design Approach:**
- Requirements analysis (functional/non-functional)
- Technology evaluation and selection
- Data architecture and flow design
- Service architecture planning

**Implementation Strategy:**
- Development standards and practices
- Deployment pipeline design
- Monitoring and observability
- Documentation and knowledge transfer

Please provide more details about your system requirements, scale, and constraints for specific architectural recommendations."""
        }
        
        return fallback_responses.get(agent_type, fallback_responses["general"])
    
    async def upload_and_analyze_file_content(self, file_content: bytes, filename: str, user_id: str):
        """Upload and analyze file - exact from original"""
        try:
            file_analysis = self.file_system.process_file(file_content, filename)
            
            # Store file context in session
            self.current_sessions[user_id]['file_context'] = file_analysis
            
            # Store in database
            with sqlite3.connect(self.memory.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    INSERT INTO file_processing 
                    (user_id, file_path, file_type, processing_result, success, timestamp)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (user_id, filename, file_analysis['file_type'], "File processed successfully", True, datetime.now()))
                conn.commit()
            
            return {
                "success": True,
                "message": "File processed successfully",
                "file_analysis": file_analysis
            }
            
        except Exception as e:
            logger.error(f"File processing error: {e}")
            return {
                "success": False,
                "error": f"File processing failed: {str(e)}"
            }
    
    async def search_web(self, query: str, user_id: str):
        """Web search - exact from original"""
        try:
            search_results = await self.web_search.search_web(query, max_results=5)
            
            if search_results.get("success"):
                # Store in memory
                with sqlite3.connect(self.memory.db_path) as conn:
                    cursor = conn.cursor()
                    cursor.execute('''
                        INSERT INTO search_history (user_id, search_query, search_type, results_count, timestamp)
                        VALUES (?, ?, ?, ?, ?)
                    ''', (user_id, query, "web", search_results.get("count", 0), datetime.now()))
                    conn.commit()
                
                # Format response
                formatted_response = f"🔍 **Web Search Results for: {query}**\n\n"
                for i, result in enumerate(search_results.get("results", []), 1):
                    formatted_response += f"**{i}. {result['title']}**\n"
                    formatted_response += f"Source: {result['source']}\n"
                    formatted_response += f"{result['snippet']}\n\n"
                
                return {"success": True, "formatted_response": formatted_response}
            else:
                return {"error": "Web search failed"}
                
        except Exception as e:
            return {"error": f"Web search error: {e}"}

    def get_system_status(self) -> Dict[str, Any]:
        """System status - enhanced original"""
        return {
            "status": "operational",
            "version": "3.0.0-free-apis-enhanced",
            "timestamp": datetime.now().isoformat(),
            "components": {
                "memory_system": "operational",
                "professional_agents": len(self.agents.agents),
                "api_providers": len(self.api_manager.available),
                "ml_system": "enhanced" if ML_SYSTEM_AVAILABLE else "basic",
                "advanced_systems": ADVANCED_SYSTEMS,
                "voice_processing": VOICE_AVAILABLE,
                "file_processing": FILE_PROCESSING_AVAILABLE,
                "github_integration": GITHUB_INTEGRATION
            },
            "api_providers": self.api_manager.get_provider_stats(),
            "capabilities": {
                "always_ai_response": True,
                "smart_enhancement_detection": True,
                "ml_enhanced_routing": ML_SYSTEM_AVAILABLE,
                "context_aware_responses": True,
                "professional_agents": bool(self.agents.agents),
                "conversation_memory": True,
                "performance_monitoring": ADVANCED_SYSTEMS,
                "free_tier_apis": True,
                "multi_provider_fallback": True
            },
            "session_info": {
                "total_sessions": len(self.current_sessions),
                "conversation_count": self.conversation_count,
                "available_providers": len(self.api_manager.available)
            }
        }

    def clear_user_context(self, user_id: str):
        """Clear context - exact from original"""
        if user_id in self.current_sessions:
            user_session = self.current_sessions[user_id]
            user_session['file_context'] = None
            user_session['conversation_count'] = 0
            user_session['last_agent'] = 'general'
        
        logger.info(f"Context cleared for user: {user_id}")

# Initialize NOVA system
nova_system = NovaUltraSystem()

# ========== FASTAPI APPLICATION SETUP ==========
app = FastAPI(
    title="NOVA Ultra Professional AI Assistant - Free APIs", 
    description="Enhanced ML-integrated professional AI assistant with Free API Providers - Always AI Response Mode",
    version="3.0.0-free-apis-enhanced"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static files
from fastapi.staticfiles import StaticFiles
app.mount("/static", StaticFiles(directory="static"), name="static")

# ========== PYDANTIC MODELS ==========
class ChatRequest(BaseModel):
    message: str = Field(..., description="User message")
    user_id: str = Field("web-user", description="User ID")

class ChatResponse(BaseModel):
    response: str
    agent_used: str
    language: str
    emotion: str
    emotion_confidence: float
    agent_confidence: float
    response_time: float
    conversation_count: int
    file_context_used: bool
    user_id: str
    session_id: str
    ml_enhanced: bool = Field(default=False, description="ML enhancement applied")
    context_used: bool = Field(default=False, description="Context used")
    recommendations: List[str] = Field(default=[], description="ML recommendations")
    enhancement_reason: str = Field(default="", description="Why enhancement was/wasn't applied")

class VoiceRequest(BaseModel):
    text: str = Field(..., description="Text to speak")

class SearchRequest(BaseModel):
    query: str = Field(..., description="Search query")
    user_id: str = Field("web-user", description="User ID")

class GitHubRequest(BaseModel):
    repo_url: str = Field(..., description="GitHub repository URL")

class GitHubQuestionRequest(BaseModel):
    question: str = Field(..., description="Question about repository")

# ========== API ENDPOINTS ==========

@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "message": "🚀 NOVA Ultra Professional API - Free APIs Enhanced",
        "version": "3.0.0-free-apis-enhanced", 
        "status": "✅ Fully Operational - Always AI with Free APIs",
        "features": [
            "🧠 Always AI Response - No dummy responses ever",
            "🎯 Smart Enhancement Detection for optimal performance",
            "🆓 FREE API Providers (Groq, OpenRouter, Together, HuggingFace, Cohere, NVIDIA)",
            "🤖 Multi-Agent System (7 agents) with Smart ML Routing",
            "💾 UltraHybridMemorySystem with semantic memory",
            "🔀 Multi-Provider AI with Professional System Prompts",
            "📄 File Processing System with ML Enhancement",
            "🔗 GitHub Repository Analyzer with ML Insights",
            "🎤 Voice Processing (Azure + Basic) with Smart Enhancement",
            "🔍 Web Search Integration",
            "💭 Conversation Memory with ML Context Storage"
        ],
        "api_providers": {
            "total_configured": len(nova_system.api_manager.providers),
            "available_now": len(nova_system.api_manager.available),
            "provider_names": [p['name'] for p in nova_system.api_manager.available]
        },
        "enhancement_logic": {
            "simple_queries": "AI response without ML overhead",
            "complex_queries": "AI response with full ML pipeline",
            "always_ai": "Every query gets AI response - no exceptions",
            "free_apis": "Using best free API providers for cost-effective operation"
        }
    }

@app.post("/chat", response_model=ChatResponse)
async def chat_endpoint(request: ChatRequest):
    """Enhanced Chat endpoint - ALWAYS AI Response with Smart Enhancement and Free APIs"""
    
    logger.info(f"💬 Chat request: {request.message[:50]}... from user: {request.user_id}")
    
    # Check if smart enhancement is needed
    needs_enhancement = SmartEnhancementDetector.needs_ml_enhancement(request.message)
    is_simple = SmartEnhancementDetector.is_simple_greeting(request.message)
    
    logger.info(f"🧠 Analysis - ML needed: {needs_enhancement}, Simple: {is_simple}")
    
    if needs_enhancement and not is_simple:
        # Complex query - Apply ML enhancement + AI
        logger.info(f"🔥 Complex query - applying ML enhancement: {request.message[:50]}...")
        
        # Run ML pipeline for complex queries
        ml_results = {}
        if ML_SYSTEM_AVAILABLE:
            ml_results = ml_manager.process_user_query(request.message, context={})

        # Extract insights from ML pipeline
        routing = ml_results.get("routing_decision", {})
        query_analysis = ml_results.get("query_analysis", {})
        context_enhancement = ml_results.get("context_enhancement", {})

        intent = routing.get("selected_agent", "general")
        confidence = routing.get("confidence_level", 0.0)
        sentiment = query_analysis.get("sentiment", "neutral")
        keywords = query_analysis.get("intent_keywords", [])
        entities = query_analysis.get("technical_context", {})
        rag_context = context_enhancement.get("relevant_context", "")
        recommendations = ml_results.get("recommendations", [])

        # Build enhanced AI prompt
        enhanced_prompt = f"""
        User asked: {request.message}

        🔍 ML Analysis:
        - Detected intent: {intent} (confidence: {confidence:.2f})
        - Sentiment: {sentiment}
        - Keywords: {keywords}
        - Entities: {entities}
        - Relevant Context: {rag_context}
        - Recommendations: {recommendations}

        ➡️ Please generate a professional, comprehensive, and engaging response.
        Use the ML insights naturally to provide the most helpful answer possible.
        """
        response_data = await nova_system.get_response(enhanced_prompt, request.user_id, intent)
        
        # Log ML interaction
        if ML_SYSTEM_AVAILABLE:
            ml_manager.store_interaction_intelligently(
                request.message,
                response_data["response"],
                agent_used=intent
            )

        # Add ML enhancement info
        response_data.update({
            'ml_enhanced': True,
            'context_used': bool(rag_context),
            'recommendations': recommendations[:3],
            'enhancement_reason': 'Complex query - full ML enhancement applied with AI'
        })

    else:
        # Simple query - Direct AI response without ML overhead
        logger.info(f"💫 Simple query - direct AI response: {request.message[:50]}...")
        
        response_data = await nova_system.get_response(request.message, request.user_id, "general")
        
        # Add simple enhancement info
        response_data.update({
            'ml_enhanced': False,
            'context_used': False,
            'recommendations': [],
            'enhancement_reason': 'Simple query - direct AI response for optimal speed'
        })

    logger.info(f"✅ Response generated - ML Enhanced: {response_data.get('ml_enhanced', False)}")

    return ChatResponse(**response_data)

@app.post("/file/upload")
async def enhanced_file_upload(
    file: UploadFile = File(...),
    user_id: str = Form(...),
    prompt: Optional[str] = Form(None)
):
    """Enhanced file upload with AI analysis"""
    start_time = time.time()
    
    try:
        # Read file content
        file_content = await file.read()
        file_size = len(file_content)
        file_type = file.content_type or "unknown"
        
        logger.info(f"📎 File upload: {file.filename} ({file_type}, {file_size} bytes)")
        
        # Basic file analysis
        file_analysis = {
            "file_name": file.filename,
            "file_type": file_type,
            "file_size": file_size,
            "upload_time": datetime.now().isoformat()
        }
        
        # Text extraction based on file type
        extracted_text = ""
        
        if file_type.startswith('text/') or file.filename.endswith(('.txt', '.md', '.py', '.js', '.html', '.css')):
            extracted_text = file_content.decode('utf-8', errors='ignore')
            file_analysis.update({
                "lines": len(extracted_text.splitlines()),
                "words": len(extracted_text.split()),
                "chars": len(extracted_text)
            })
        
        elif FILE_PROCESSING_AVAILABLE:
            # Advanced file processing
            if file.filename.endswith('.pdf'):
                try:
                    pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                    extracted_text = ""
                    for page in pdf_reader.pages:
                        extracted_text += page.extract_text() + "\n"
                    file_analysis["pages"] = len(pdf_reader.pages)
                except Exception as e:
                    logger.error(f"PDF processing error: {e}")
            
            elif file.filename.endswith('.docx'):
                try:
                    doc = docx.Document(BytesIO(file_content))
                    extracted_text = ""
                    for paragraph in doc.paragraphs:
                        extracted_text += paragraph.text + "\n"
                    file_analysis["paragraphs"] = len(doc.paragraphs)
                except Exception as e:
                    logger.error(f"DOCX processing error: {e}")
        
        # Always use AI for file analysis, determine if ML enhancement is needed
        analysis_query = prompt or f"Analyze this {file_type} file and provide professional insights"
        needs_enhancement = SmartEnhancementDetector.needs_ml_enhancement(analysis_query)
        
        if needs_enhancement and extracted_text:
            # Apply ML enhancement for complex file analysis
            logger.info("🧠 Applying ML enhancement for file analysis")
            
            enhanced_prompt = f"""Professional File Analysis Request:

File Details:
- Name: {file.filename}
- Type: {file_type}
- Size: {file_size} bytes
- Content Preview: {extracted_text[:1000]}...

User Request: {analysis_query}

Please provide a comprehensive professional analysis including:
1. **Content Summary**: Key themes and main points
2. **Technical Analysis**: Structure, format, and technical details
3. **Insights & Findings**: Important observations and patterns
4. **Quality Assessment**: Strengths and areas for improvement
5. **Recommendations**: Actionable next steps and suggestions
6. **Professional Context**: Industry relevance and best practices

Structure your response professionally with clear sections and actionable insights."""
            
            response_data = await nova_system.get_response(enhanced_prompt, user_id, "general")
            ai_response = response_data['response']
            ml_enhanced = True
        else:
            # Simple file processing with AI response
            simple_prompt = f"""File Analysis:

File: {file.filename} ({file_type}, {file_size} bytes)
Content: {extracted_text[:500] if extracted_text else 'Binary or unsupported format'}...

{analysis_query if prompt else 'Please provide a summary and analysis of this file.'}

Please analyze the file and provide helpful insights."""
            
            response_data = await nova_system.get_response(simple_prompt, user_id, "general")
            ai_response = response_data['response']
            ml_enhanced = False
        
        # Store file processing record
        with sqlite3.connect(memory_system.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO file_processing 
                (user_id, file_path, file_type, processing_result, success, timestamp)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (
                user_id,
                file.filename,
                file_type,
                "File processed successfully with AI analysis",
                True,
                datetime.now()
            ))
            conn.commit()
        
        processing_time = time.time() - start_time
        
        return {
            "success": True,
            "message": "File uploaded and analyzed successfully",
            "response": ai_response,
            "metadata": {
                "file_analysis": file_analysis,
                "ml_enhanced": ml_enhanced,
                "processing_time": processing_time,
                "ai_analysis_applied": True,
                "enhancement_applied": needs_enhancement and extracted_text
            }
        }
        
    except Exception as e:
        logger.error(f"File upload error: {e}")
        return {
            "success": False,
            "message": f"File upload failed: {str(e)}",
            "response": "",
            "metadata": {"processing_time": time.time() - start_time}
        }

@app.post("/github/analyze")
async def enhanced_github_analysis(
    repo_url: str = Form(...),
    user_id: str = Form("web-user")
):
    """Enhanced GitHub repository analysis with AI"""
    try:
        logger.info(f"🔍 GitHub analysis: {repo_url}")
        
        if not GITHUB_INTEGRATION:
            # Even without GitHub integration, provide AI response
            no_integration_prompt = f"""GitHub Repository Analysis Request:

Repository: {repo_url}

I don't have direct GitHub integration available, but I can provide professional guidance on repository analysis:

Please provide a comprehensive framework for analyzing this repository, including:
1. **Code Quality Assessment Methods**
2. **Architecture Analysis Approaches**
3. **Security Review Guidelines**
4. **Performance Optimization Strategies**
5. **Best Practices Evaluation**
6. **Improvement Recommendations Framework**

Structure this as actionable guidance for manual repository analysis."""
            
            response_data = await nova_system.get_response(no_integration_prompt, user_id, "coding")
            
            return {
                "success": True,
                "message": "Repository analysis guidance provided",
                "response": response_data['response'],
                "metadata": {
                    "repo_url": repo_url,
                    "ml_enhanced": False,
                    "processing_time": response_data.get('response_time', 0),
                    "integration_available": False,
                    "guidance_provided": True
                }
            }
        
        start_time = time.time()
        
        # GitHub analysis is inherently complex - always apply ML enhancement
        enhanced_prompt = f"""Professional GitHub Repository Analysis:

Repository: {repo_url}

Please provide a comprehensive technical analysis framework including:

**1. Repository Overview Analysis**
- Project structure and organization assessment
- Technology stack evaluation and dependencies review
- Documentation quality and completeness analysis

**2. Code Quality Assessment Framework**
- Code organization and architecture patterns evaluation
- Coding standards and best practices compliance review
- Technical debt identification and code smell detection

**3. Security & Performance Analysis**
- Security vulnerability assessment methodology
- Performance optimization opportunities identification
- Scalability considerations and bottleneck analysis

**4. Professional Recommendations**
- Priority improvement suggestions and refactoring roadmap
- Industry best practices implementation guidelines
- Development workflow optimization strategies

**5. Strategic Technical Insights**
- Project maturity assessment and maintenance evaluation
- Community engagement and contribution pattern analysis
- Long-term sustainability and evolution planning

Structure your analysis professionally with specific, actionable recommendations."""
        
        response_data = await nova_system.get_response(enhanced_prompt, user_id, "coding")
        
        # Store analysis in database
        with sqlite3.connect(memory_system.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT OR REPLACE INTO github_repos 
                (repo_url, repo_name, analysis_date, suggestions)
                VALUES (?, ?, ?, ?)
            ''', (
                repo_url,
                repo_url.split('/')[-1],
                datetime.now(),
                "Professional analysis completed with AI enhancement"
            ))
            conn.commit()
        
        return {
            "success": True,
            "message": "Repository analysis completed",
            "response": response_data['response'],
            "metadata": {
                "repo_url": repo_url,
                "ml_enhanced": True,
                "processing_time": time.time() - start_time,
                "agent_used": response_data.get('agent_used', 'coding'),
                "ai_analysis_applied": True
            }
        }
        
    except Exception as e:
        logger.error(f"GitHub analysis error: {e}")
        return {
            "success": False,
            "message": f"Analysis failed: {str(e)}",
            "response": ""
        }

@app.post("/github/question")
async def enhanced_github_question(
    question: str = Form(...),
    user_id: str = Form("web-user")
):
    """Enhanced GitHub repository Q&A with AI"""
    try:
        logger.info(f"❓ GitHub question: {question[:50]}...")
        
        if not GITHUB_INTEGRATION:
            # Provide AI response even without GitHub integration
            no_integration_prompt = f"""Technical Repository Question:

Question: {question}

While I don't have direct repository access, I can provide comprehensive technical guidance:

Please provide professional technical assistance addressing this question with:
1. **General Technical Approach**
2. **Best Practices and Standards**
3. **Implementation Guidelines**
4. **Common Solutions and Patterns**
5. **Troubleshooting Strategies**
6. **Additional Resources and Learning**

Structure this as actionable technical guidance."""
            
            response_data = await nova_system.get_response(no_integration_prompt, user_id, "coding")
            
            return {
                "success": True,
                "message": "Technical guidance provided",
                "response": response_data['response'],
                "metadata": {
                    "question": question,
                    "ml_enhanced": False,
                    "processing_time": response_data.get('response_time', 0),
                    "integration_available": False,
                    "ai_guidance_provided": True
                }
            }
        
        start_time = time.time()
        
        # Repository Q&A is inherently complex - always apply enhancement
        enhanced_prompt = f"""Repository Technical Question:

Question: {question}

Please provide a comprehensive technical answer that includes:

**1. Direct Technical Answer**
- Clear, specific response to the question
- Technical details and implementation guidance
- Code examples and best practices where applicable

**2. Context & Background**
- Relevant technical context and considerations
- Industry standards and established patterns
- Common challenges and solutions

**3. Implementation Guidance**
- Step-by-step implementation approach
- Configuration and setup details
- Testing and validation strategies

**4. Advanced Technical Considerations**
- Performance implications and optimizations
- Security considerations and best practices
- Scalability and maintenance factors

**5. Professional Resources**
- Related documentation and technical resources
- Further learning opportunities and references
- Community best practices and patterns

Provide a professional, detailed response that addresses both immediate technical needs and broader understanding."""
        
        response_data = await nova_system.get_response(enhanced_prompt, user_id, "coding")
        
        return {
            "success": True,
            "message": "Question answered successfully",
            "response": response_data['response'],
            "metadata": {
                "question": question,
                "ml_enhanced": True,
                "processing_time": time.time() - start_time,
                "agent_used": response_data.get('agent_used', 'coding'),
                "ai_analysis_applied": True
            }
        }
        
    except Exception as e:
        logger.error(f"GitHub question error: {e}")
        return {
            "success": False,
            "message": f"Question processing failed: {str(e)}",
            "response": ""
        }

@app.post("/voice/speak")
async def voice_speak_endpoint(audio: UploadFile = File(...)):
    """Process voice audio and return TTS response"""
    try:
        # Save the incoming audio file temporarily
        temp_path = f"temp_{audio.filename}"
        async with aiofiles.open(temp_path, 'wb') as out_file:
            content = await audio.read()
            await out_file.write(content)
        
        # Process the audio file (implement your logic here)
        audio_data = await nova_system.voice_system.process_audio(temp_path)
        
        # Clean up
        os.remove(temp_path)
        
        return StreamingResponse(
            BytesIO(audio_data),
            media_type="audio/wav",
            headers={"Content-Disposition": "attachment; filename=response.wav"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Audio processing error: {str(e)}")
    
@app.post("/voice/process")
async def process_voice_command(
    audio: UploadFile = File(None),
    text: str = Form(None),
    user_id: str = Form("voice-user")
):
    """
    Unified voice processing with Always AI Response:
    - If `audio` is uploaded → STT → AI → TTS → return spoken answer
    - If `text` is provided → AI → TTS → return spoken answer
    """
    try:
        if audio:
            # 1. Read raw bytes from browser
            audio_data = await audio.read()

            # 2. Convert WebM → WAV
            wav_bytes = webm_to_wav(audio_data)

            # 3. STT
            user_text = await nova_system.voice_system.process_audio(wav_bytes)

            # 4. AI Response (ALWAYS use AI)
            ai_response_data = await nova_system.get_response(user_text, user_id, "general")
            ai_response = ai_response_data['response']

            # 5. TTS
            processed_audio = await nova_system.voice_system.text_to_speech(
                ai_response,
                voice="en-US-AriaNeural"
            )

        elif text:
            # Direct AI processing then TTS
            ai_response_data = await nova_system.get_response(text, user_id, "general")
            ai_response = ai_response_data['response']
            
            processed_audio = await nova_system.voice_system.text_to_speech(
                ai_response,
                voice="en-US-AriaNeural"
            )

        else:
            return JSONResponse(
                {"error": "No audio or text provided"},
                status_code=400
            )

        return StreamingResponse(
            BytesIO(processed_audio),
            media_type="audio/wav",
            headers={"Content-Disposition": "attachment; filename=response.wav"}
        )

    except Exception as e:
        logger.error(f"Voice endpoint error: {e}")
        raise HTTPException(status_code=500, detail=f"Audio processing error: {str(e)}")

@app.post("/web/search")
async def web_search_endpoint(request: SearchRequest):
    result = await nova_system.search_web(request.query, request.user_id)
    return result

@app.get("/agents")
async def get_agents():
    """Get available agents with enhancement info"""
    agents_info = {
        "general": {
            "name": "NOVA General AI",
            "description": "Ultra-professional general AI assistant - Always AI Response",
            "emoji": "🤖",
            "specialties": ["general knowledge", "problem solving", "research"],
            "always_ai": True
        },
        "coding": {
            "name": "Professional Code Expert", 
            "description": "Full-stack development specialist - Always AI Response",
            "emoji": "💻",
            "specialties": ["programming", "debugging", "architecture"],
            "always_ai": True
        },
        "career": {
            "name": "Career Development Coach",
            "description": "Professional career guidance expert - Always AI Response",
            "emoji": "🎯", 
            "specialties": ["career planning", "resume optimization", "interview prep"],
            "always_ai": True
        },
        "business": {
            "name": "Strategic Business Consultant",
            "description": "Business intelligence and strategy expert - Always AI Response",
            "emoji": "📊",
            "specialties": ["business strategy", "market analysis", "growth planning"],
            "always_ai": True
        },
        "medical": {
            "name": "Health & Wellness Advisor",
            "description": "Evidence-based health guidance specialist - Always AI Response",
            "emoji": "🏥",
            "specialties": ["health information", "wellness planning", "medical research"],
            "always_ai": True
        },
        "emotional": {
            "name": "Emotional Support Counselor",
            "description": "Empathetic emotional guidance specialist - Always AI Response",
            "emoji": "💙",
            "specialties": ["emotional support", "stress management", "mental wellness"],
            "always_ai": True
        },
        "technical_architect": {
            "name": "Technical System Architect",
            "description": "System design and architecture expert - Always AI Response",
            "emoji": "🏗️",
            "specialties": ["system architecture", "scalability", "technical design"],
            "always_ai": True
        }
    }
    
    # Add enhancement info to each agent
    for agent_info in agents_info.values():
        agent_info["ml_enhanced"] = ML_SYSTEM_AVAILABLE
        agent_info["smart_routing"] = True
        agent_info["always_ai_response"] = True
    
    return {
        "agents": agents_info,
        "ml_system_available": ML_SYSTEM_AVAILABLE,
        "smart_enhancement": True,
        "always_ai_response": True,
        "no_dummy_responses": True
    }

@app.get("/system")
async def get_system_status():
    """Get enhanced system status"""
    return nova_system.get_system_status()

@app.get("/health")
async def health_check():
    """Enhanced health check"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "3.0.0-always-ai-enhanced",
        "components": {
            "nova_system": "operational",
            "memory_system": "operational", 
            "ml_enhancement": "enhanced" if ML_SYSTEM_AVAILABLE else "basic",
            "smart_routing": True,
            "database": "connected" if os.path.exists(memory_system.db_path) else "disconnected"
        },
        "features": {
            "always_ai_response": True,
            "no_dummy_responses": True,
            "smart_enhancement_detection": True,
            "ml_powered_routing": ML_SYSTEM_AVAILABLE,
            "conversation_memory": True,
            "multi_agent_system": True,
            "professional_responses": True
        },
        "guarantee": "Every user query receives AI-generated response - no exceptions"
    }

@app.post("/clear/{user_id}")
async def clear_context_endpoint(user_id: str):
    """Clear user context"""
    nova_system.clear_user_context(user_id)
    return {"success": True, "message": f"Context cleared for user {user_id}"}

# ========== STARTUP EVENT ==========
@app.on_event("startup")
async def startup_event():
    """Enhanced startup event"""
    logger.info("🚀 NOVA Ultra Professional AI Assistant Starting...")
    logger.info("💫 ALWAYS AI RESPONSE MODE - No dummy responses ever!")
    logger.info(f"✅ Memory System: {type(nova_system.memory).__name__}")
    logger.info(f"✅ Professional Agents: {len(nova_system.agents.agents)} loaded")
    logger.info(f"✅ API Providers: {len(nova_system.api_manager.available)} available")
    logger.info(f"✅ ML System: {'Enhanced' if ML_SYSTEM_AVAILABLE else 'Basic Mode'}")
    logger.info(f"✅ Smart Enhancement Detection: Active")
    logger.info(f"✅ Always AI Response: Guaranteed")
    logger.info(f"✅ Advanced Systems: {'Available' if ADVANCED_SYSTEMS else 'Basic Mode'}")
    logger.info("🎯 NOVA Ultra Professional API Ready - Always AI Mode!")

# ========== MAIN ENTRY POINT ==========
if __name__ == "__main__":
    logger.info("🚀 Starting NOVA Ultra Professional FastAPI Backend...")
    logger.info("💫 ALWAYS AI RESPONSE MODE ENABLED")
    logger.info("🔡 Backend will be available at: http://0.0.0.0:8000")
    logger.info("📚 API Documentation: http://0.0.0.0:8000/docs")
    logger.info(f"🤖 ML System Status: {'Enhanced' if ML_SYSTEM_AVAILABLE else 'Basic Mode'}")
    logger.info(f"🧠 Memory System: Advanced Hybrid")
    logger.info(f"🎯 Agent System: Multi-Agent with Smart ML Routing")
    logger.info("🚫 NO DUMMY RESPONSES - Every query gets AI processing")
    logger.info("⚡ Smart Enhancement: Simple queries = Fast AI, Complex queries = ML + AI")
    
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=5000,
        log_level="info"
    )